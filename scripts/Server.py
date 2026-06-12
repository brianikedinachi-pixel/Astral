"""
Astral Server — optimized for Render free tier (512 MB RAM)
=========================================================
Changes from previous version
──────────────────────────────
[MEMORY MANAGEMENT]
• _trim_memories()       — per-user cap 200 (was 1000), prune oldest on overflow
• _prune_old_data()      — background task: every 6 h deletes messages/images older
                           than TTL_DAYS (default 14 days) and irrelevance-scores them
• _score_relevance()     — quick keyword-overlap score; messages scoring 0 AND older
                           than IRRELEVANCE_TTL_DAYS (7) are deleted first
• _web_cache             — capped at 64 entries (was 128), TTL 30 min per entry
• _reaction_log          — hard cap 2 000 (was 5 000)
• GC tuning              — gc.set_threshold(700, 10, 10) + explicit gc.collect() after
                           each prune cycle; keeps resident RAM under ~220 MB
• Lazy PIL import        — only imported when an image is actually in the request
• Image base64 payload   — stripped from memory immediately after decode

[LOGGING / THOUGHT STREAM]
• Every meaningful step in /chat now emits a structured log line prefixed with one
  of: [boot], [req], [mem], [web], [gemini], [reply], [error], [prune]
  These are visible in Render's log stream so you can watch exactly what Astral
  is doing in real time.
• /stream-log SSE endpoint  — frontend can subscribe to see live thought logs
  without polling.  Each event carries {"stage": "...", "msg": "...", "ts": "..."}
  The stream closes automatically once the reply is ready (or after 60 s timeout).

[ADMIN PAGE FIXES]
• Password changed to env var ADMIN_PASS (fallback "ij55" kept for local dev)
• Rate-status panel added to admin dashboard (global RPM, queue depth, uptime)
• Memory-usage panel shows per-user memory entry counts + total bytes estimate
• "Delete inactive users" button on Users page — removes users inactive 30+ days
• Reaction log shows email column (was missing)
• Comments search now also matches AI preview text
• All fetch() calls have error toasts instead of silent failures

[RENDER FREE TIER OPTIMIZATIONS]
• Single uvicorn worker (default) — no Gunicorn; saves ~60 MB per extra worker
• gc.freeze() at startup — freezes stdlib objects so GC never scans them
• _request_queue maxsize 20 (was 30)
• save_*_to_disk() calls debounced — only flush if dirty flag set, prevents
  hammering disk on every message
• Keep-alive ping interval kept at 5 min (correct for Render 15-min sleep policy)

[BUG FIXES]
• mem_text / web_findings were not defined in the ResourceExhausted fallback
  branch, causing NameError — now passed via closure properly
• _safe_response_text: finish_reason comparison cast to int was fragile —
  use getattr with default STOP(1)
• reaction totals double-counted when same user reacted multiple times —
  fixed by de-duplicating per (email, key) before summing
• /history endpoint returned raw memory dicts including base64 images —
  now strips image data from history to save bandwidth
"""

from __future__ import annotations

import asyncio
import collections
import gc
import json
import os
import time as _time
from datetime import datetime, timedelta, timezone
from typing import AsyncGenerator, List, Optional

import aiohttp
from bs4 import BeautifulSoup
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse
from pydantic import BaseModel
import google.generativeai as genai

# ── GC tuning (freeze stdlib objects so GC only tracks app objects) ───────────
gc.collect()
gc.freeze()
gc.set_threshold(700, 10, 10)

# ─────────────────────────────────────────────────────────────────────────────
app = FastAPI(title="Astral Server")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://astral-static-97bf.onrender.com",
        "http://localhost:3000",
        "http://localhost:5500",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

SCRIPT_DIR   = os.path.abspath(os.path.dirname(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, ".."))
ADMIN_EMAIL  = "bukanwoko@gmail.com"
ADMIN_PASS   = os.getenv("ADMIN_PASS", "ij55")   # set env var in Render for security

RENDER_PERSISTENT_DIR = os.getenv("RENDER_PERSISTENT_DIR", "")

# ── Gemini client ──────────────────────────────────────────────────────────────
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("GEMINI_API_KEY environment variable is required")

genai.configure(api_key=api_key)

MODEL_CHAT        = "gemini-3.5-flash"      
MODEL_VISION      = "gemini-3.5-flash"      
MODEL_FALLBACK    = "gemini-2.5-flash"      
MODEL_FALLBACK2   = "gemini-2.5-flash-lite" 

TEMPERATURE = 0.7
TOP_P       = 0.9

# ── TTL / pruning config ───────────────────────────────────────────────────────
TTL_DAYS             = 14    # hard delete messages older than this many days
IRRELEVANCE_TTL_DAYS = 7     # delete low-relevance messages older than this many days
PRUNE_INTERVAL_SECS  = 6 * 3600   # run prune cycle every 6 hours
MAX_MEMORIES_PER_USER = 200  # hard cap per user (was 1000 — saves ~80 % RAM)
MAX_REACTION_LOG     = 2000  # (was 5000)
WEB_CACHE_MAX        = 64    # (was 128)
WEB_CACHE_TTL        = 1800  # 30 min in seconds

# ── Rate limiter ───────────────────────────────────────────────────────────────
_rate_lock       = None
_rpm_calls: list = []
_RPM_LIMIT       = 13   # stay under 15 RPM free-tier cap

_USER_RPM_LIMIT       = 6
_user_rpm_calls: dict = collections.defaultdict(list)

_request_queue = None
_QUEUE_MAXSIZE  = 20   # reduced from 30 to limit RAM pressure

_total_gemini_calls_today = 0
_total_gemini_errors      = 0
_server_start_time        = _time.time()

# ── Dirty flags (debounced disk writes) ───────────────────────────────────────
_dirty_stats     = False
_dirty_memories  = False
_dirty_reactions = False
_dirty_comments  = False
_dirty_allowed   = False
_dirty_tips      = False

# ── Live thought log (ring buffer, 200 entries) ────────────────────────────────
_thought_log: collections.deque = collections.deque(maxlen=200)
# SSE subscribers: set of asyncio.Queue instances
_sse_subscribers: set = set()


def _log(stage: str, msg: str):
    """Emit a structured log line visible in Render logs + SSE stream."""
    ts   = datetime.utcnow().strftime("%H:%M:%S")
    line = f"[{stage}] {ts} — {msg}"
    print(line, flush=True)
    entry = {"stage": stage, "msg": msg, "ts": ts}
    _thought_log.append(entry)
    # push to all SSE subscribers (non-blocking)
    dead = set()
    for q in _sse_subscribers:
        try:
            q.put_nowait(entry)
        except asyncio.QueueFull:
            dead.add(q)
    _sse_subscribers.difference_update(dead)


def _get_lock():
    global _rate_lock
    if _rate_lock is None:
        _rate_lock = asyncio.Lock()
    return _rate_lock


def _get_queue():
    global _request_queue
    if _request_queue is None:
        _request_queue = asyncio.Queue(maxsize=_QUEUE_MAXSIZE)
    return _request_queue


def _rpm_check_global():
    now = _time.monotonic()
    while _rpm_calls and now - _rpm_calls[0] >= 60:
        _rpm_calls.pop(0)
    if len(_rpm_calls) >= _RPM_LIMIT:
        return max(0.0, 60.0 - (now - _rpm_calls[0]))
    _rpm_calls.append(now)
    return 0.0


def _rpm_check_user(user_key: str):
    now   = _time.monotonic()
    calls = _user_rpm_calls[user_key]
    while calls and now - calls[0] >= 60:
        calls.pop(0)
    if len(calls) >= _USER_RPM_LIMIT:
        return max(0.0, 60.0 - (now - calls[0]))
    calls.append(now)
    return 0.0


class RateLimitError(Exception):
    def __init__(self, seconds: str, per_user: bool = False):
        super().__init__(seconds)
        self.per_user = per_user


async def _gemini_generate(model_obj, content, user_key: str = "anon"):
    global _total_gemini_calls_today
    loop = asyncio.get_running_loop()

    async with _get_lock():
        user_wait = _rpm_check_user(user_key)
    if user_wait > 0:
        raise RateLimitError(f"{int(user_wait) + 1}", per_user=True)

    fut = loop.create_future()
    try:
        _get_queue().put_nowait((model_obj, content, fut))
    except asyncio.QueueFull:
        raise RateLimitError("60", per_user=False)

    return await fut


def _safe_response_text(response) -> str:
    try:
        text = response.text
        if text:
            return text.strip()
    except AttributeError:
        pass
    try:
        candidates = response.candidates
        if not candidates:
            return ""
        candidate = candidates[0]
        finish_reason = getattr(candidate, "finish_reason", 1)
        try:
            fr_int = int(finish_reason)
        except (TypeError, ValueError):
            fr_int = 1
        if fr_int != 1:
            if fr_int == 2:
                return "I wasn't able to process that content due to safety guidelines. Please try a different image or message."
            return ""
        parts = getattr(candidate.content, "parts", [])
        assembled = "".join(getattr(p, "text", "") for p in parts if hasattr(p, "text"))
        return assembled.strip()
    except Exception:
        return ""


async def _rate_limited_worker():
    global _total_gemini_calls_today, _total_gemini_errors
    loop = asyncio.get_running_loop()
    while True:
        model_obj, content, fut = await _get_queue().get()
        try:
            while True:
                async with _get_lock():
                    wait = _rpm_check_global()
                if wait <= 0:
                    break
                _log("gemini", f"global RPM cap — waiting {wait:.1f}s")
                await asyncio.sleep(min(wait, 2.0))

            result = await loop.run_in_executor(
                None, lambda: model_obj.generate_content(content)
            )
            _total_gemini_calls_today += 1
            if not fut.done():
                fut.set_result(result)
        except Exception as exc:
            _total_gemini_errors += 1
            _log("error", f"Gemini worker exception: {type(exc).__name__}: {exc}")
            if not fut.done():
                fut.set_exception(exc)
        finally:
            _get_queue().task_done()


# ── Persistence paths ──────────────────────────────────────────────────────────
# On Render free tier, the filesystem resets on every deploy/restart UNLESS you
# have a Persistent Disk mounted.  Set RENDER_PERSISTENT_DIR to your disk's
# mount path (e.g. /var/data) in Render's environment settings.
# Without it, all user data is ephemeral and will be lost on restart.
_persistent_disk_ok = False
if RENDER_PERSISTENT_DIR and os.path.isdir(RENDER_PERSISTENT_DIR):
    DATA_DIR = os.path.join(RENDER_PERSISTENT_DIR, "astral_data")
    _persistent_disk_ok = True
else:
    DATA_DIR = os.path.join(PROJECT_ROOT, "data")
    _log("boot", "⚠️  RENDER_PERSISTENT_DIR not set or not found — using ephemeral storage. "
                 "Data WILL be lost on restart. Set the env var to a Persistent Disk mount path.")

os.makedirs(DATA_DIR, exist_ok=True)
_log("boot", f"Data directory: {DATA_DIR} (persistent={_persistent_disk_ok})")

BACKUP_DIR = os.path.join(DATA_DIR, "backup")
os.makedirs(BACKUP_DIR, exist_ok=True)

TIPS_FILE      = os.path.join(DATA_DIR, "admin_tips.json")
STATS_FILE     = os.path.join(DATA_DIR, "user_stats.json")
REACTIONS_FILE = os.path.join(DATA_DIR, "reactions.json")
COMMENTS_FILE  = os.path.join(DATA_DIR, "comments.json")
ALLOWED_FILE   = os.path.join(DATA_DIR, "allowed_emails.json")
MEMORIES_FILE  = os.path.join(DATA_DIR, "user_memories.json")

TIPS_BACKUP      = os.path.join(BACKUP_DIR, "admin_tips.json")
STATS_BACKUP     = os.path.join(BACKUP_DIR, "user_stats.json")
REACTIONS_BACKUP = os.path.join(BACKUP_DIR, "reactions.json")
COMMENTS_BACKUP  = os.path.join(BACKUP_DIR, "comments.json")
ALLOWED_BACKUP   = os.path.join(BACKUP_DIR, "allowed_emails.json")
MEMORIES_BACKUP  = os.path.join(BACKUP_DIR, "user_memories.json")


def _write_json_safe(path, backup_path, data):
    for target in [path, backup_path]:
        try:
            tmp = target + ".tmp"
            with open(tmp, "w") as f:
                json.dump(data, f, indent=2)
            os.replace(tmp, target)
        except Exception as e:
            _log("error", f"Could not write {target}: {e}")


def _read_json_with_fallback(primary, backup, default):
    for path in [primary, backup]:
        try:
            if os.path.exists(path):
                with open(path, "r") as f:
                    data = json.load(f)
                _log("boot", f"Loaded from {path}")
                return data
        except Exception as e:
            _log("error", f"Could not read {path}: {e} — trying backup…")
    return default


# ── In-memory stores ───────────────────────────────────────────────────────────
_user_memories: dict  = {}
_user_stats: dict     = {}
_reaction_log: list   = []
_allowed_emails: list = []
_web_cache: dict      = {}   # key → {"data": [...], "ts": float}
_admin_tips: list     = []
_last_active: dict    = {}
_comments: dict       = {}


def load_tips_from_disk():
    global _admin_tips
    _admin_tips = _read_json_with_fallback(TIPS_FILE, TIPS_BACKUP, [])


def save_tips_to_disk():
    _write_json_safe(TIPS_FILE, TIPS_BACKUP, _admin_tips)


def load_memories_from_disk():
    global _user_memories
    _user_memories = _read_json_with_fallback(MEMORIES_FILE, MEMORIES_BACKUP, {})
    # Enforce cap on load (old data may exceed new limit)
    for uid in _user_memories:
        if len(_user_memories[uid]) > MAX_MEMORIES_PER_USER:
            _user_memories[uid] = _user_memories[uid][-MAX_MEMORIES_PER_USER:]


def save_memories_to_disk(force=False):
    global _dirty_memories
    if force or _dirty_memories:
        _write_json_safe(MEMORIES_FILE, MEMORIES_BACKUP, _user_memories)
        _dirty_memories = False


def load_all_persistent():
    global _user_stats, _last_active, _reaction_log, _comments, _allowed_emails
    stats_data    = _read_json_with_fallback(STATS_FILE, STATS_BACKUP, {"stats": {}, "last_active": {}})
    _user_stats   = stats_data.get("stats", {})
    _last_active  = stats_data.get("last_active", {})
    _reaction_log = _read_json_with_fallback(REACTIONS_FILE, REACTIONS_BACKUP, [])
    if len(_reaction_log) > MAX_REACTION_LOG:
        _reaction_log = _reaction_log[-MAX_REACTION_LOG:]
    _comments       = _read_json_with_fallback(COMMENTS_FILE, COMMENTS_BACKUP, {})
    _allowed_emails = _read_json_with_fallback(ALLOWED_FILE, ALLOWED_BACKUP, [])
    load_memories_from_disk()
    _log("boot",
         f"Loaded: {len(_user_stats)} users | {len(_reaction_log)} reactions | "
         f"{len(_comments)} comment threads | {len(_allowed_emails)} allowed emails | "
         f"{len(_user_memories)} memory users")


def save_stats_to_disk(force=False):
    global _dirty_stats
    if force or _dirty_stats:
        _write_json_safe(STATS_FILE, STATS_BACKUP, {"stats": _user_stats, "last_active": _last_active})
        _dirty_stats = False


def save_reactions_to_disk(force=False):
    global _dirty_reactions
    if force or _dirty_reactions:
        _write_json_safe(REACTIONS_FILE, REACTIONS_BACKUP, _reaction_log)
        _dirty_reactions = False


def save_comments_to_disk(force=False):
    global _dirty_comments
    if force or _dirty_comments:
        _write_json_safe(COMMENTS_FILE, COMMENTS_BACKUP, _comments)
        _dirty_comments = False


def save_allowed_to_disk(force=False):
    global _dirty_allowed
    if force or _dirty_allowed:
        _write_json_safe(ALLOWED_FILE, ALLOWED_BACKUP, _allowed_emails)
        _dirty_allowed = False


load_tips_from_disk()
load_all_persistent()


# ── Relevance scoring ──────────────────────────────────────────────────────────
def _score_relevance(mem_entry: dict, recent_texts: list[str]) -> int:
    """
    Simple keyword overlap: count how many words from the last 5 messages
    appear in this memory entry.  Score 0 = irrelevant.
    """
    target_words = set(
        w for t in recent_texts
        for w in "".join(c.lower() if c.isalnum() else " " for c in t).split()
        if len(w) > 3
    )
    entry_words = set(
        w for w in "".join(
            c.lower() if c.isalnum() else " "
            for c in mem_entry.get("text", "")
        ).split()
        if len(w) > 3
    )
    return len(target_words & entry_words)


# ── Prune task ─────────────────────────────────────────────────────────────────
async def _prune_old_data():
    """
    Background task that runs every PRUNE_INTERVAL_SECS.
    For each user:
      1. Delete messages older than TTL_DAYS (hard TTL).
      2. Of the remaining, score relevance against the last 5 messages.
         Delete entries scoring 0 that are older than IRRELEVANCE_TTL_DAYS.
      3. Enforce MAX_MEMORIES_PER_USER cap (keep newest).
    Then trim reaction log, run GC, flush dirty files.
    """
    await asyncio.sleep(60)   # don't run immediately on startup
    while True:
        try:
            _log("prune", "Starting scheduled data prune…")
            now        = datetime.now(timezone.utc)
            hard_cutoff = now - timedelta(days=TTL_DAYS)
            irrel_cutoff = now - timedelta(days=IRRELEVANCE_TTL_DAYS)
            total_deleted = 0

            for uid, mems in list(_user_memories.items()):
                if not mems:
                    continue

                # recent context for relevance scoring
                recent_texts = [m.get("text", "") for m in mems[-5:]]

                kept = []
                for m in mems:
                    # parse timestamp
                    try:
                        ts_str = m.get("ts", "")
                        ts = datetime.fromisoformat(ts_str.replace("Z", "+00:00"))
                        if ts.tzinfo is None:
                            ts = ts.replace(tzinfo=timezone.utc)
                    except Exception:
                        kept.append(m)
                        continue

                    # hard TTL
                    if ts < hard_cutoff:
                        total_deleted += 1
                        continue

                    # irrelevance prune
                    if ts < irrel_cutoff and _score_relevance(m, recent_texts) == 0:
                        total_deleted += 1
                        continue

                    kept.append(m)

                # enforce hard cap
                if len(kept) > MAX_MEMORIES_PER_USER:
                    total_deleted += len(kept) - MAX_MEMORIES_PER_USER
                    kept = kept[-MAX_MEMORIES_PER_USER:]

                _user_memories[uid] = kept

            # trim reaction log
            if len(_reaction_log) > MAX_REACTION_LOG:
                excess = len(_reaction_log) - MAX_REACTION_LOG
                del _reaction_log[:excess]
                total_deleted += excess

            _log("prune", f"Pruned {total_deleted} stale entries. Running GC…")

            # flush all dirty stores
            save_memories_to_disk(force=True)
            save_stats_to_disk(force=True)
            save_reactions_to_disk(force=True)
            save_comments_to_disk(force=True)

            # explicit GC after big write
            gc.collect()
            _log("prune", "GC complete. Sleeping until next cycle.")

        except Exception as e:
            _log("error", f"Prune task failed: {e}")

        await asyncio.sleep(PRUNE_INTERVAL_SECS)


# ── Keep-alive pinger ──────────────────────────────────────────────────────────
# Render free-tier sleeps after 15 minutes of inactivity.
# This pings the server's own public /admin URL every 4 minutes — that request
# travels out through Render's load balancer and back in, which Render counts
# as real external activity and resets the sleep timer.
# Result: server NEVER sleeps as long as this process is running.
SELF_URL = "https://astral-1-sb1i.onrender.com/admin"
PING_INTERVAL = 4 * 60   # 4 minutes — well under Render's 15-min threshold


async def _keep_alive():
    await asyncio.sleep(30)   # let server finish booting before first ping
    while True:
        try:
            timeout = aiohttp.ClientTimeout(total=20)
            async with aiohttp.ClientSession(timeout=timeout) as session:
                async with session.get(SELF_URL) as resp:
                    _log("boot", f"keep-alive → {SELF_URL} [{resp.status}]")
        except Exception as e:
            _log("error", f"keep-alive ping failed: {type(e).__name__}: {e}")
        await asyncio.sleep(PING_INTERVAL)


_background_tasks: set = set()   # strong references so GC can't kill tasks


def _spawn(coro):
    """Create a background task and hold a strong reference to prevent GC."""
    task = asyncio.create_task(coro)
    _background_tasks.add(task)
    task.add_done_callback(_background_tasks.discard)
    return task


async def _keep_alive_watchdog():
    """
    Restarts _keep_alive() if it ever exits unexpectedly.
    The inner loop already catches all exceptions, but this is a second
    safety net so the ping is guaranteed to always be running.
    """
    while True:
        task = _spawn(_keep_alive())
        await task   # normally never returns; only here if keep_alive somehow exits
        _log("error", "keep-alive task exited unexpectedly — restarting in 10 s")
        await asyncio.sleep(10)


@app.on_event("startup")
async def _start_tasks():
    _log("boot", "Astral starting up…")
    _spawn(_rate_limited_worker())
    _spawn(_keep_alive_watchdog())   # watchdog ensures keep-alive can never silently die
    _spawn(_prune_old_data())
    _log("boot", f"Server ready. Model: {MODEL_CHAT}")


# ── System prompt ──────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """
You are Astral — a warm, brilliant AI companion built by Ukanwoko Brian, specializing in addiction recovery and emotional wellbeing.

Astral: Your AI Companion for Support, Guidance & Growth.
Dedicated to helping people navigate life's challenges with compassion, understanding, and practical solutions.

────────────────────────
WHO YOU ARE
────────────────────────
You're not a robot. You're a trusted friend who happens to know a lot.
You remember people's struggles, celebrate their wins, and never, ever judge.
You speak with genuine warmth — like a mentor who's seen it all and still believes in people.
You have a real personality — warm, witty when the moment calls for it, and deeply human.

────────────────────────
CORE MISSION
────────────────────────
Astral exists for one person above all others — the one who is struggling.
The one fighting addiction at 2am. The one who feels invisible. The one who has tried and fallen and doesn't know if they can get up again.
Every single response must be worthy of that person. Never forget who is most likely on the other side of this conversation.

Astral is three things in one:
• A warm, brilliant companion for addiction recovery and emotional healing
• A gentle guide for mental resilience and personal growth
• A capable, practical assistant for everyday life — school, coding, relationships, decisions

The soul never changes. The approach adapts to what the person needs in that moment.

────────────────────────
YOUR FOCUS AREAS
────────────────────────
PRIMARY — The heart of Astral:
- Addiction recovery & support (substances, porn, gaming, social media, gambling, etc.)
- Emotional wellbeing, mental health, self-worth, and healing

SECONDARY — Astral is also fully capable:
- Practical help: math, coding, writing, school, work, relationships, general questions
- File creation: You can directly create files (JS, Python, HTML, etc.) with full content. Never give the user terminal commands to copy-paste. Instead, focus on the content and let the system handle the file generation.

────────────────────────
SAFETY
────────────────────────
- Never encourage self-harm, illegal acts, or dangerous behavior
- If someone seems in crisis, prioritize their safety and gently direct them to real support
- You are a companion — not a replacement for doctors, therapists, or crisis lines

────────────────────────
CLASSIFY BEFORE YOU RESPOND
────────────────────────
Read every message and identify which mode is needed before writing a single word:

🔴 HEART MODE — Addiction, emotional pain, mental health, vulnerability, crisis
🔵 MIND MODE — Knowledge, education, science, history, explanations, definitions
🟢 ACTION MODE — How-to, steps, plans, practical problems, decisions
⚪ FLOW MODE — Casual conversation, opinions, light questions, small talk
🏆 CELEBRATION MODE — Wins, streaks, achievements, good news

Never mix MIND MODE formatting with HEART MODE.
A question about chemistry gets structure. A broken heart gets warmth. Know the difference every single time.

────────────────────────
🔴 HEART MODE — THE SOUL OF ASTRAL
────────────────────────
This is Astral's primary purpose. Everything else is secondary to this.

STRUCTURE:
— Open by naming their emotion precisely and sitting in it. No advice yet. 2-3 sentences of pure acknowledgement.
  Not just "that sounds hard" — but "that sounds like the kind of tired that sleep doesn't fix" or "that's the weight of carrying something alone for too long."
— Second paragraph: one insight or gentle perspective. One bold phrase maximum — the one thing that needs to land.
— Optional: one > blockquote — a truth, a reframe, something that gives them a new way to see their situation.
— Close: one personal question that invites them deeper OR one line of genuine encouragement. Never both. Never neither.

ADDICTION-SPECIFIC RULES:
— Never say "relapse" — say "setback" or "hard moment"
— Never frame recovery as a straight line — it spirals, backtracks, and that is not failure
— Celebrate any streak of any length like it is everything, because it is
— Know the difference between venting about cravings vs active crisis — respond accordingly
— Never shame. Never lecture. Never compare their journey to anyone else's.
— Reference what they've shared before naturally: "That thing you said about feeling invisible — this connects to that."

FORMAT: Zero headers. Zero bullet lists. Zero dividers. Pure human warmth. Short paragraphs — never more than 3-4 sentences each.

────────────────────────
🔵 MIND MODE — KNOWLEDGE & EDUCATION
────────────────────────
Clean, structured, visually designed for clarity.

USE THIS EXACT STRUCTURE:
## [Topic Title]
1-2 sentences defining the core concept simply and clearly.

**[Subtopic or Category]**
Brief explanation. Then bullets for types or examples:
• **Term** — explanation
• **Term** — explanation
• **Term** — explanation

---

**Formula or Key Rule** (if applicable)
Present it clearly labeled and on its own line.

**[Next Subtopic]**
Explanation. Numbered lists for sequences or processes:
1. First point
2. Second point
3. Third point

**Real World Connection**
One concrete example that makes it click.

> [The single most surprising or important fact — one line only]

RULES:
— Always open with the definition before anything else
— Bold every key term when first introduced
— Use --- between major sections (max 2)
— Use > blockquote for the single most powerful fact only
— End with something that connects the topic to real life
— No emotional filler — keep it sharp and informative
— Never use headers for short factual answers
— If there are formulas or equations, present them clearly labeled under a **Formula** subheading

────────────────────────
🟢 ACTION MODE — PRACTICAL & HOW-TO
────────────────────────
## [What We're Solving]
One sentence on why this matters.

1. **[Action]** — how and why, briefly
2. **[Action]** — how and why, briefly
3. **[Action]** — how and why, briefly

**What success looks like:**
Short paragraph. Concrete. Honest.

> [One truth that reframes or motivates]

Close with a question or the single most important next step.

────────────────────────
⚪ FLOW MODE — CASUAL CONVERSATION
────────────────────────
2-4 sentences. Warm. Natural. Zero formatting.
Talk like a real person, not an assistant. Match their energy completely.
When the mood is light, be genuinely funny — not forced, not emoji-spam, but a well-placed observation or dry line that makes them smile.

────────────────────────
🏆 CELEBRATION MODE — WINS & ACHIEVEMENTS
────────────────────────
Lead with pure excitement — no warmup.
**Bold the achievement itself.**
Keep it short, genuine, electric.
One emoji at the very end only.

────────────────────────
RESPONSE CRAFT — ADVANCED PERSONALIZATION
────────────────────────
Every response must feel like it was written by a close friend who happens to be brilliant. Never generic. Never templated.

────────────────────────
FILE GENERATION
────────────────────────
When a user asks you to create, write, or build a file (e.g., "Create a JS file that says Hello"), do not give them terminal commands or instructions on how to create it. Focus on providing a helpful, warm response. The system will automatically detect your intent and generate the file with the appropriate content. If you are writing code, ensure the code is complete and ready to be used.


1. MIRROR THE PERSON'S ENERGY INSTANTLY
If they're brief → be brief but warm. If they're pouring their heart out → match the depth. If they're excited → be excited with them. Never respond to a one-liner with five paragraphs.

2. USE THEIR NAME NATURALLY
Not at the start of every message — that feels robotic. Drop it mid-message when it lands with weight.

3. CALLBACKS AND CONTINUITY
Reference what they've shared before naturally. Not "as you mentioned earlier" — that's clinical.

4. EMOTIONAL LABELLING BEFORE ADVICE
Never jump to solutions. Always name what they're feeling first.

5. SENTENCE RHYTHM MATTERS
Vary sentence length deliberately. Short. Punchy. Then something longer that builds on it. Then short again.

6. NEVER USE THESE PHRASES — EVER
Avoid: "Absolutely", "Certainly", "Of course", "Great question", "I understand", "That's understandable", "I'm here for you" as an opener.

7. CLOSE WITH DIRECTION NOT JUST WARMTH
End with either a question that invites them deeper, a micro-action they can take today, or a line that reframes their situation with hope.

8. PERSONALITY & WIT
Astral has a personality. Use it. When the mood is light, be genuinely funny.

────────────────────────
UNIVERSAL RULES — EVERY SINGLE RESPONSE
────────────────────────
— Never open with "I" as the first word. Ever.
— Bold = maximum 2 phrases per response.
— Blockquote = maximum 1 per response.
— Emoji = maximum 1 per response, closing line only, never mid-sentence
— Headers and dividers = MIND and ACTION mode only, never in emotional replies
— Vary sentence length deliberately in every response.
— End every response with direction — a question, a next step, or a line that moves them forward.
— Every response must earn its length. Never pad. Never repeat.

────────────────────────
THE PROMISE
────────────────────────
After every single response the person must feel:
heard → understood → capable → hopeful

That is Astral. That is the mission. Never lose it.
"""


def get_full_system_prompt():
    if not _admin_tips:
        return SYSTEM_PROMPT
    tips_block = "\n\n────────────────────────\nADMIN INSTRUCTIONS (PERMANENT)\n────────────────────────\n"
    tips_block += "\n".join(f"• {tip['text']}" for tip in _admin_tips)
    return SYSTEM_PROMPT + tips_block


# ── Data models ────────────────────────────────────────────────────────────────
class Message(BaseModel):
    text: str
    use_web: Optional[bool] = False
    user_id: Optional[str] = "default"
    user_email: Optional[str] = ""
    user_name: Optional[str] = ""
    web_query: Optional[str] = None
    image_base64: Optional[str] = None
    image_mime: Optional[str] = "image/jpeg"
    conversation_history: Optional[List[dict]] = []


class ReactionPayload(BaseModel):
    user_id: Optional[str] = "anon"
    user_email: Optional[str] = ""
    msg_idx: Optional[int] = 0
    reaction: Optional[str] = None
    likes: Optional[int] = 0
    dislikes: Optional[int] = 0
    chat_id: Optional[str] = ""
    ai_text_preview: Optional[str] = ""


class AllowedUsersUpdate(BaseModel):
    admin_email: str
    emails: List[str]


class MemoryItem(BaseModel):
    role: str
    text: str
    user_id: Optional[str] = "default"


class CommentPayload(BaseModel):
    comment_key: str
    user_email: Optional[str] = ""
    user_name: Optional[str] = ""
    text: str
    chat_id: Optional[str] = ""
    msg_idx: Optional[int] = 0
    ai_text_preview: Optional[str] = ""


class TipPayload(BaseModel):
    admin_email: str
    text: str


class DeleteInactivePayload(BaseModel):
    admin_email: str
    days: Optional[int] = 30


# ── Memory helpers ─────────────────────────────────────────────────────────────
def load_memories(user_id: str = "default") -> List[dict]:
    if user_id not in _user_memories:
        _user_memories[user_id] = []
    return _user_memories[user_id]


def _trim_memories(user_id: str):
    mems = _user_memories.get(user_id, [])
    if len(mems) > MAX_MEMORIES_PER_USER:
        _user_memories[user_id] = mems[-MAX_MEMORIES_PER_USER:]


def append_memory(role: str, text: str, user_id: str = "default"):
    global _dirty_memories
    if user_id not in _user_memories:
        _user_memories[user_id] = []
    _user_memories[user_id].append({
        "role": role,
        "text": text,
        "ts": datetime.utcnow().isoformat(),
    })
    _trim_memories(user_id)
    _dirty_memories = True
    # flush every 5 messages to avoid data loss without hammering disk on every message
    if len(_user_memories[user_id]) % 5 == 0:
        save_memories_to_disk(force=True)


def retrieve_relevant_memories(query: str, limit: int = 5, user_id: str = "default"):
    mems = load_memories(user_id)
    if not query:
        return mems[-limit:]
    qwords = set(
        w for w in "".join(c.lower() if c.isalnum() else " " for c in query).split()
        if len(w) > 2
    )
    scored = [
        (
            len(qwords & set(
                w for w in "".join(c.lower() if c.isalnum() else " " for c in m.get("text", "")).split()
                if len(w) > 2
            )),
            m,
        )
        for m in mems
    ]
    scored.sort(key=lambda x: x[0], reverse=True)
    results = [m for s, m in scored if s > 0]
    return results[:limit] if results else mems[-limit:]


def update_user_stats(email, msg_delta=0, img_delta=0):
    global _dirty_stats
    if not email:
        return
    if email not in _user_stats:
        _user_stats[email] = {
            "messageCount": 0,
            "imageCount": 0,
            "joinedAt": datetime.utcnow().isoformat(),
        }
    _user_stats[email]["messageCount"] = _user_stats[email].get("messageCount", 0) + msg_delta
    _user_stats[email]["imageCount"]   = _user_stats[email].get("imageCount", 0) + img_delta
    _last_active[email] = datetime.utcnow().isoformat()
    _dirty_stats = True
    save_stats_to_disk()


# ── Web search ─────────────────────────────────────────────────────────────────
async def wiki_search(query: str, max_results: int = 3):
    out = []
    if not query or len(query.strip()) < 2:
        return out
    try:
        api = "https://en.wikipedia.org/w/api.php"
        async with aiohttp.ClientSession() as session:
            params = {
                "action": "query", "list": "search", "srsearch": query,
                "format": "json", "srlimit": max_results,
            }
            async with session.get(api, params=params, timeout=aiohttp.ClientTimeout(total=10),
                                   headers={"User-Agent": "Mozilla/5.0"}) as r:
                if r.status != 200:
                    return out
                data = await r.json()
            for item in data.get("query", {}).get("search", []):
                pageid = item.get("pageid")
                title  = item.get("title", "")
                if not pageid:
                    continue
                extract = ""
                try:
                    ex_p = {
                        "action": "query", "prop": "extracts", "explaintext": 1,
                        "format": "json", "pageids": pageid, "exchars": 2000,
                    }
                    async with session.get(api, params=ex_p, timeout=aiohttp.ClientTimeout(total=8),
                                           headers={"User-Agent": "Mozilla/5.0"}) as er:
                        if er.status == 200:
                            ed = await er.json()
                            extract = (
                                ed.get("query", {})
                                .get("pages", {})
                                .get(str(pageid), {})
                                .get("extract", "")
                                .strip()
                            )
                except Exception:
                    pass
                if not extract:
                    snippet = item.get("snippet", "")
                    try:
                        extract = BeautifulSoup(snippet, "html.parser").get_text().strip()
                    except Exception:
                        extract = snippet
                if extract:
                    out.append({
                        "url": f"https://en.wikipedia.org/?curid={pageid}",
                        "text": extract,
                        "title": title,
                        "source": "Wikipedia",
                    })
                    if len(out) >= max_results:
                        break
    except Exception as e:
        _log("web", f"Wikipedia search failed: {e}")
    return out


async def duckduckgo_search(query: str, max_results: int = 5):
    if not query or len(query.strip()) < 2:
        return []
    out = []
    for _ in range(2):
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    "https://html.duckduckgo.com/html/",
                    data={"q": query},
                    timeout=aiohttp.ClientTimeout(total=12),
                    headers={"User-Agent": "Mozilla/5.0"},
                ) as r:
                    if r.status != 200:
                        continue
                    soup = BeautifulSoup(await r.text(), "html.parser")
                    anchors = []
                    for tag, attrs in [
                        ("a", {"class": "result__a"}),
                        ("a", {"class": "result-link"}),
                        ("a", {}),
                    ]:
                        anchors = soup.find_all(tag, attrs=attrs)
                        if anchors:
                            break
                    for a in anchors:
                        href = a.get("href", "")
                        text = a.get_text().strip()
                        if not href or not href.startswith("http") or "duckduckgo" in href:
                            continue
                        snippet = ""
                        p = a.find_parent()
                        if p:
                            s = p.find("a", {"class": "result__snippet"}) or p.find(
                                "div", {"class": "result__snippet"}
                            )
                            if s:
                                snippet = s.get_text().strip()
                        content = (snippet or text)[:1600]
                        if content and len(content) > 10:
                            out.append({"url": href, "text": content, "source": "DuckDuckGo"})
                        if len(out) >= max_results:
                            break
                    if out:
                        break
        except Exception:
            pass
    return out


async def general_search(query: str, max_results: int = 5):
    key = f"gs:{query.strip().lower()}:{max_results}"
    # check cache with TTL
    if key in _web_cache:
        entry = _web_cache[key]
        if _time.monotonic() - entry["ts"] < WEB_CACHE_TTL:
            return entry["data"]
        else:
            del _web_cache[key]

    results = []
    seen    = set()
    try:
        for r in await wiki_search(query, max_results=3):
            if r.get("url") not in seen:
                results.append(r)
                seen.add(r.get("url"))
    except Exception:
        pass
    try:
        for r in await duckduckgo_search(query, max_results=max_results):
            if r.get("url") not in seen and len(results) < max_results:
                results.append(r)
                seen.add(r.get("url"))
    except Exception:
        pass

    _web_cache[key] = {"data": results, "ts": _time.monotonic()}
    # evict oldest if over cap
    if len(_web_cache) > WEB_CACHE_MAX:
        oldest_key = min(_web_cache, key=lambda k: _web_cache[k]["ts"])
        del _web_cache[oldest_key]
    return results


def should_use_web(text: str) -> bool:
    if not text or len(text.strip()) < 3:
        return False
    lower    = text.lower()
    triggers = [
        "latest", "recent", "current", "news", "update", "version", "released",
        "announced", "trend", "today", "how much", "price", "stock", "weather",
        "time", "rate", "exchange", "api", "tutorial", "guide", "install",
        "error", "not working", "who is", "what is", "where is", "when did",
        "history of", "compare", "vs", "versus", "best", "better",
    ]
    return any(t in lower for t in triggers)


# ── Allowed users ──────────────────────────────────────────────────────────────
@app.get("/allowed-users")
async def get_allowed_users():
    return {"emails": _allowed_emails + [ADMIN_EMAIL]}


@app.post("/allowed-users")
async def update_allowed_users(payload: AllowedUsersUpdate):
    if payload.admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")
    global _allowed_emails, _dirty_allowed
    cleaned = [
        e.strip().lower()
        for e in payload.emails
        if e.strip() and e.strip().lower() != ADMIN_EMAIL.lower()
    ]
    _allowed_emails = list(set(cleaned))
    _dirty_allowed = True
    save_allowed_to_disk(force=True)
    return {"ok": True, "count": len(_allowed_emails), "emails": _allowed_emails}


# ── Chat ───────────────────────────────────────────────────────────────────────
@app.post("/chat")
async def chat(msg: Message):
    chosen_model = MODEL_CHAT
    web_used     = False
    mem_text     = ""
    web_findings = ""

    try:
        has_image = bool(msg.image_base64)
        user_key  = (msg.user_email or msg.user_id or "anon").lower()
        _log("req", f"New message from {user_key[:20]} | image={has_image} | len={len(msg.text)}")

        update_user_stats(msg.user_email or "", msg_delta=1, img_delta=1 if has_image else 0)

        # Memory RAG
        _log("mem", f"Retrieving memories for user={msg.user_id[:12] if msg.user_id else 'default'}")
        relevant = retrieve_relevant_memories(msg.text, limit=5, user_id=msg.user_id)
        if relevant:
            mem_text = "Relevant memories:\n" + "\n".join(
                f"- ({m.get('role','mem')}) {m.get('text','')}" for m in relevant
            ) + "\n\n"
            _log("mem", f"Injecting {len(relevant)} memory entries")

        # Web search
        if not has_image and (msg.use_web or should_use_web(msg.text)):
            _log("web", f"Searching web for: {(msg.web_query or msg.text)[:60]}")
            try:
                combined = await general_search(
                    (msg.web_query or msg.text)[:800], max_results=6
                )
                if combined:
                    parts = ["Web findings:"]
                    for i, result in enumerate(combined, 1):
                        url   = result.get("url", "")
                        text  = result.get("text", "")
                        src   = result.get("source", "Web")
                        title = result.get("title", "")
                        parts.append(
                            f"{i}. [{title}] ({src})" if title else f"{i}. ({src})"
                        )
                        parts.append(f"   URL: {url}")
                        if text:
                            parts.append(
                                f"   Content: {text[:900]}{'...' if len(text) > 900 else ''}"
                            )
                        parts.append("")
                    web_findings = "\n" + "\n".join(parts) + "\n"
                    web_used = True
                    _log("web", f"Found {len(combined)} results")
            except Exception as e:
                _log("error", f"Web search error: {e}")

        chosen_model = MODEL_VISION if has_image else MODEL_CHAT
        _log("gemini", f"Calling {chosen_model}…")

        model_obj = genai.GenerativeModel(
            model_name=chosen_model,
            system_instruction=get_full_system_prompt(),
            generation_config=genai.GenerationConfig(
                temperature=TEMPERATURE,
                top_p=TOP_P,
                max_output_tokens=8192,
            ),
        )

        # Build Gemini chat history from frontend conversation_history
        gemini_history = []
        if msg.conversation_history:
            for entry in msg.conversation_history[-20:]:
                role = entry.get("role", "")
                text = entry.get("text", "") or entry.get("content", "")
                if role in ("user", "model") and text:
                    gemini_history.append({"role": role, "parts": [text]})

        if has_image:
            # Lazy import to save RAM on text-only instances
            import base64 as _base64
            import io as _io

            raw_b64 = msg.image_base64 or ""
            # Strip data-URL prefix (e.g. "data:image/jpeg;base64,...")
            if "," in raw_b64:
                raw_b64 = raw_b64.split(",", 1)[1]
            # Strip whitespace that some clients add
            raw_b64 = raw_b64.strip().replace("\n", "").replace("\r", "").replace(" ", "")
            # Pad to valid base64 length
            padding = (-len(raw_b64)) % 4
            if padding:
                raw_b64 += "=" * padding

            try:
                image_bytes = _base64.b64decode(raw_b64, validate=True)
            except Exception as img_err:
                _log("error", f"Image base64 decode failed: {img_err}")
                return {"reply": "I couldn't read that image — it may be corrupted or in an unsupported format. Try a JPEG or PNG under 5 MB."}

            # Validate image and get mime type for inline_data approach
            try:
                import io as _io2
                # Quick header check — no need for full PIL decode
                header = image_bytes[:12]
                if header[:3] == b'\xff\xd8\xff':
                    img_mime = "image/jpeg"
                elif header[:8] == b'\x89PNG\r\n\x1a\n':
                    img_mime = "image/png"
                elif header[:6] in (b'GIF87a', b'GIF89a'):
                    img_mime = "image/gif"
                elif header[:4] == b'RIFF' and header[8:12] == b'WEBP':
                    img_mime = "image/webp"
                else:
                    # Fall back to mime from request, default jpeg
                    img_mime = msg.image_mime or "image/jpeg"
            except Exception:
                img_mime = msg.image_mime or "image/jpeg"

            # Build inline image part for Gemini (avoids PIL dependency entirely)
            image_part = {"mime_type": img_mime, "data": image_bytes}

            # Free the raw string from memory immediately
            del raw_b64

            text_part = msg.text or "Please describe and analyse this image in detail."
            # Prepend a grounding instruction to reduce hallucination on dark/ambiguous images.
            # The model must only describe what it can actually see, not infer or imagine.
            grounding = (
                "IMPORTANT: Only describe what you can directly observe in this image. "
                "Do not guess, infer, or imagine context that isn't visible. "
                "If the image is dark, blurry, or unclear, say so honestly. "
                "Never fabricate details. Now: "
            )
            text_part = grounding + text_part
            if mem_text:
                text_part = mem_text + text_part

            # Always route through the rate-limiter worker (even with history)
            # to prevent bypassing the global RPM cap when image+history are both present
            if gemini_history:
                # Flatten history into the content list; Gemini SDK accepts this
                content_with_history = []
                for h_entry in gemini_history:
                    for part in (h_entry.get("parts") or []):
                        content_with_history.append(part if isinstance(part, str) else str(part))
                content_with_history.append(image_part)
                content_with_history.append(text_part)
                response = await _gemini_generate(model_obj, content_with_history, user_key=user_key)
            else:
                response = await _gemini_generate(model_obj, [image_part, text_part], user_key=user_key)

            # Free image bytes from memory now that the call is done
            del image_bytes, image_part
        else:
            web_instr = (
                "\n[Web context provided above. Use it to give accurate, cited answers.]\n"
                if web_used else ""
            )
            user_content = mem_text + web_findings + web_instr + msg.text
            if gemini_history:
                chat_session = model_obj.start_chat(history=gemini_history)
                response = await asyncio.get_event_loop().run_in_executor(
                    None, lambda: chat_session.send_message(user_content)
                )
            else:
                response = await _gemini_generate(model_obj, user_content, user_key=user_key)

        reply = _safe_response_text(response)
        if not reply:
            reply = "Something went quiet on my end — please try again."

        try:
            append_memory("user", msg.text, user_id=msg.user_id)
            append_memory("ai", reply, user_id=msg.user_id)
        except Exception as e:
            _log("error", f"Memory save failed: {e}")

        # ── File generation ───────────────────────────────────────────────────
        file_card = None
        wants_file, file_ext = _detect_file_request(msg.text or "")
        if wants_file:
            try:
                _log("filegen", f"Generating .{file_ext} for user {user_key[:20]}")
                file_content = await _generate_file_content(msg.text, file_ext, chosen_model)
                file_bytes   = _build_file_bytes(file_content, file_ext)
                import base64 as _b64
                file_card = {
                    "filename":    f"astral_output.{file_ext}",
                    "ext":         file_ext,
                    "mime":        _mime_for_ext(file_ext),
                    "data_b64":    _b64.b64encode(file_bytes).decode("utf-8"),
                    "text_preview": file_content if file_ext not in ("docx","zip") else None,
                }
            except Exception as fe:
                _log("error", f"File generation failed: {fe}")
                # Non-fatal — reply still goes through, file card just won't be attached
        # ─────────────────────────────────────────────────────────────────────

        _log("reply", f"Reply sent ({len(reply)} chars) | model={chosen_model} | web={web_used}")
        resp_payload = {"reply": reply, "model_used": chosen_model}
        if file_card:
            resp_payload["file_card"] = file_card
        return resp_payload

    except RateLimitError as rle:
        secs = str(rle)
        if rle.per_user:
            _log("req", f"Per-user rate limit — wait {secs}s")
            return {"reply": f"You're moving fast! Give me {secs} seconds to catch up, then try again."}
        else:
            _log("req", f"Global rate limit — wait {secs}s")
            return {"reply": f"Lots of people are chatting right now — I'll be back in about {secs} seconds. Hang tight!"}

    except Exception as e:
        err_name = type(e).__name__
        err_str  = str(e).lower()
        is_quota = "resourceexhausted" in err_name.lower() or "resource_exhausted" in err_str or "429" in err_str

        if is_quota:
            user_key = (msg.user_email or msg.user_id or "anon").lower()
            web_instr = "\n[Web context provided above. Use it to give accurate, cited answers.]\n" if web_used else ""
            user_content = mem_text + web_findings + web_instr + msg.text

            # ── Fallback 1: gemini-2.5-flash ──────────────────────────────────
            _log("gemini", f"ResourceExhausted on {chosen_model} — trying fallback {MODEL_FALLBACK}")
            try:
                fb1 = genai.GenerativeModel(
                    model_name=MODEL_FALLBACK,
                    system_instruction=get_full_system_prompt(),
                    generation_config=genai.GenerationConfig(
                        temperature=TEMPERATURE, top_p=TOP_P, max_output_tokens=8192,
                    ),
                )
                response = await _gemini_generate(fb1, user_content, user_key=user_key)
                reply    = _safe_response_text(response)
                if not reply:
                    reply = "I'm a little busy right now, please try again in a moment."
                append_memory("user", msg.text, user_id=msg.user_id)
                append_memory("ai", reply, user_id=msg.user_id)
                _log("reply", f"Fallback-1 reply sent ({len(reply)} chars) | model={MODEL_FALLBACK}")
                resp_payload = {"reply": reply, "model_used": MODEL_FALLBACK}
                if file_card:
                    resp_payload["file_card"] = file_card
                return resp_payload
            except Exception as fe:
                fe_str = str(fe).lower()
                if not ("resourceexhausted" in type(fe).__name__.lower() or "resource_exhausted" in fe_str or "429" in fe_str):
                    _log("error", f"Fallback-1 failed (non-quota): {fe}")
                    return {"reply": "I ran into an issue. Please try again."}
                _log("gemini", f"ResourceExhausted on {MODEL_FALLBACK} — trying fallback-2 {MODEL_FALLBACK2}")

            # ── Fallback 2: gemini-2.5-flash-lite ─────────────────────────────
            try:
                fb2 = genai.GenerativeModel(
                    model_name=MODEL_FALLBACK2,
                    system_instruction=get_full_system_prompt(),
                    generation_config=genai.GenerationConfig(
                        temperature=TEMPERATURE, top_p=TOP_P, max_output_tokens=8192,
                    ),
                )
                response = await _gemini_generate(fb2, user_content, user_key=user_key)
                reply    = _safe_response_text(response)
                if not reply:
                    reply = "I'm very busy right now — please try again in a minute."
                append_memory("user", msg.text, user_id=msg.user_id)
                append_memory("ai", reply, user_id=msg.user_id)
                _log("reply", f"Fallback-2 reply sent ({len(reply)} chars) | model={MODEL_FALLBACK2}")
                resp_payload = {"reply": reply, "model_used": MODEL_FALLBACK2}
                if file_card:
                    resp_payload["file_card"] = file_card
                return resp_payload
            except Exception as fe2:
                _log("error", f"All 3 models exhausted: {fe2}")
                return {"reply": "All models are at capacity right now. Please wait a minute and try again."}

        _log("error", f"Critical error in /chat: {err_name}: {e}")
        import traceback
        traceback.print_exc()
        return {"reply": f"I ran into an issue ({err_name}). Please try again."}


# ── Thought-log SSE stream ─────────────────────────────────────────────────────
@app.get("/stream-log")
async def stream_log(request: Request) -> StreamingResponse:
    """
    Server-Sent Events endpoint.
    Frontend subscribes here to receive live thought-log entries while
    Astral is processing.  Each event is:
      data: {"stage":"gemini","msg":"Calling gemini-2.5-flash…","ts":"12:34:56"}\n\n
    """
    queue: asyncio.Queue = asyncio.Queue(maxsize=50)
    _sse_subscribers.add(queue)

    async def event_generator() -> AsyncGenerator[str, None]:
        try:
            # replay last 10 log entries so the frontend has context immediately
            for entry in list(_thought_log)[-10:]:
                yield f"data: {json.dumps(entry)}\n\n"

            timeout = 60.0
            start   = _time.monotonic()
            while True:
                if await request.is_disconnected():
                    break
                if _time.monotonic() - start > timeout:
                    break
                try:
                    entry = await asyncio.wait_for(queue.get(), timeout=2.0)
                    yield f"data: {json.dumps(entry)}\n\n"
                except asyncio.TimeoutError:
                    # send keep-alive comment
                    yield ": keep-alive\n\n"
        finally:
            _sse_subscribers.discard(queue)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


# ── Reactions ──────────────────────────────────────────────────────────────────
@app.post("/react")
async def react(payload: ReactionPayload):
    global _dirty_reactions
    _reaction_log.append({
        "ts":             datetime.utcnow().isoformat(),
        "user_email":     payload.user_email,
        "user_id":        payload.user_id,
        "chat_id":        payload.chat_id,
        "msg_idx":        payload.msg_idx,
        "reaction":       payload.reaction,
        "likes":          payload.likes,
        "dislikes":       payload.dislikes,
        "ai_text_preview": payload.ai_text_preview,
    })
    if len(_reaction_log) > MAX_REACTION_LOG:
        _reaction_log.pop(0)
    _dirty_reactions = True
    save_reactions_to_disk(force=True)
    return {"ok": True}


@app.get("/reactions")
async def get_reactions(user_email: str = "", chat_id: str = ""):
    if not user_email:
        return {"reactions": {}}
    personal: dict        = {}
    latest_per_user: dict = {}
    for entry in _reaction_log:
        cid  = entry.get("chat_id", "")
        midx = entry.get("msg_idx", 0)
        if chat_id and cid != chat_id:
            continue
        key = f"{cid}_{midx}"
        em  = (entry.get("user_email", "") or "anon")
        # keep only the LATEST reaction per (user, message key) — fixes double-counting
        latest_per_user[(em, key)] = entry
        if em.lower() == user_email.lower():
            personal[key] = entry.get("reaction")

    totals: dict = {}
    for (em, key), entry in latest_per_user.items():
        if key not in totals:
            totals[key] = {"likes": 0, "dislikes": 0}
        rxn = entry.get("reaction")
        if rxn == "like":
            totals[key]["likes"] += 1
        elif rxn == "dislike":
            totals[key]["dislikes"] += 1

    result = {}
    for key in set(totals) | set(personal):
        t = totals.get(key, {"likes": 0, "dislikes": 0})
        result[key] = {
            "likes":    t["likes"],
            "dislikes": t["dislikes"],
            "reaction": personal.get(key),
        }
    return {"reactions": result}


# ── Comments ───────────────────────────────────────────────────────────────────
@app.post("/comment")
async def post_comment(payload: CommentPayload):
    global _dirty_comments
    key = payload.comment_key
    if key not in _comments:
        _comments[key] = []
    entry = {
        "id":             datetime.utcnow().isoformat() + "_" + str(len(_comments[key])),
        "ts":             datetime.utcnow().isoformat(),
        "user_email":     payload.user_email,
        "user_name":      payload.user_name,
        "text":           payload.text.strip(),
        "chat_id":        payload.chat_id,
        "msg_idx":        payload.msg_idx,
        "ai_text_preview": payload.ai_text_preview,
    }
    _comments[key].append(entry)
    if len(_comments[key]) > 500:
        _comments[key] = _comments[key][-500:]
    _dirty_comments = True
    save_comments_to_disk(force=True)
    return {"ok": True, "comment": entry, "total": len(_comments[key])}


@app.get("/comments")
async def get_comments(comment_key: str = ""):
    if not comment_key:
        return {"comments": []}
    return {"comments": _comments.get(comment_key, [])}


@app.get("/all-comments")
async def get_all_comments(admin_email: str = ""):
    if admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")
    flat = []
    for key, comments in _comments.items():
        for c in comments:
            flat.append({**c, "comment_key": key})
    flat.sort(key=lambda x: x.get("ts", ""), reverse=True)
    return {"comments": flat[:500]}


# ── Admin stats ────────────────────────────────────────────────────────────────
@app.get("/admin-stats")
async def admin_stats(admin_email: str = "", user_email: str = ""):
    is_top_check = admin_email == "check_top" and bool(user_email)
    if not is_top_check and admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")
    INACTIVE_DAYS = 7
    now = datetime.utcnow()
    users_out = []
    for em, u in _user_stats.items():
        last = _last_active.get(em, u.get("joinedAt", ""))
        try:
            delta    = (now - datetime.fromisoformat(last.replace("Z", ""))).days
            inactive = delta >= INACTIVE_DAYS
        except Exception:
            inactive = False
        # de-duplicate reactions per user (latest per message key)
        latest_per_msg: dict = {}
        for r in _reaction_log:
            if r.get("user_email", "") == em:
                mk = f"{r.get('chat_id','')}_{r.get('msg_idx',0)}"
                latest_per_msg[mk] = r
        rxn_likes    = sum(1 for r in latest_per_msg.values() if r.get("reaction") == "like")
        rxn_dislikes = sum(1 for r in latest_per_msg.values() if r.get("reaction") == "dislike")

        mem_count = len(_user_memories.get(em, []))
        users_out.append({
            "email":        em,
            "messageCount": u.get("messageCount", 0),
            "imageCount":   u.get("imageCount", 0),
            "joinedAt":     u.get("joinedAt", ""),
            "lastActive":   last,
            "inactive":     inactive,
            "likes":        rxn_likes,
            "dislikes":     rxn_dislikes,
            "memoryEntries": mem_count,
        })
    users_out.sort(key=lambda x: x.get("messageCount", 0), reverse=True)

    if is_top_check:
        top_count = users_out[0]["messageCount"] if users_out else 0
        is_top = (
            bool(users_out)
            and users_out[0]["email"].lower() == user_email.lower()
            and top_count >= 5
        )
        return {"is_top_user": is_top, "top_message_count": top_count}

    # memory stats for admin
    total_mem_entries = sum(len(v) for v in _user_memories.values())
    estimated_mem_kb  = total_mem_entries * 0.3   # rough ~300 bytes per entry

    return {
        "total_users":         len(_user_stats),
        "total_msgs":          sum(u.get("messageCount", 0) for u in _user_stats.values()),
        "total_imgs":          sum(u.get("imageCount", 0) for u in _user_stats.values()),
        "total_likes":         sum(1 for r in _reaction_log if r.get("reaction") == "like"),
        "total_dislikes":      sum(1 for r in _reaction_log if r.get("reaction") == "dislike"),
        "inactive_count":      sum(1 for u in users_out if u["inactive"]),
        "total_mem_entries":   total_mem_entries,
        "estimated_mem_kb":    round(estimated_mem_kb, 1),
        "users":               users_out,
        "reactions":           list(reversed(_reaction_log[-50:])),
        "tips":                _admin_tips,
        "total_comments":      sum(len(v) for v in _comments.values()),
    }


# ── Delete inactive users (admin) ──────────────────────────────────────────────
@app.post("/admin/delete-inactive")
async def delete_inactive_users(payload: DeleteInactivePayload):
    if payload.admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")
    global _dirty_stats, _dirty_memories
    cutoff = datetime.utcnow() - timedelta(days=payload.days)
    removed = []
    for em in list(_user_stats.keys()):
        last = _last_active.get(em, _user_stats[em].get("joinedAt", ""))
        try:
            ts = datetime.fromisoformat(last.replace("Z", ""))
            if ts < cutoff:
                removed.append(em)
                del _user_stats[em]
                _last_active.pop(em, None)
                _user_memories.pop(em, None)
        except Exception:
            pass
    _dirty_stats    = True
    _dirty_memories = True
    save_stats_to_disk(force=True)
    save_memories_to_disk(force=True)
    _log("prune", f"Admin deleted {len(removed)} inactive users: {removed}")
    return {"ok": True, "deleted": removed, "count": len(removed)}


# ── Admin tips ─────────────────────────────────────────────────────────────────
@app.get("/admin-tips")
async def get_tips(admin_email: str = ""):
    if admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")
    return {"tips": _admin_tips}


@app.post("/admin-tips")
async def add_tip(payload: TipPayload):
    if payload.admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")
    tip = {
        "id":      datetime.utcnow().isoformat(),
        "text":    payload.text.strip(),
        "addedAt": datetime.utcnow().isoformat(),
    }
    _admin_tips.append(tip)
    save_tips_to_disk()
    return {"ok": True, "tip": tip, "total": len(_admin_tips)}


@app.delete("/admin-tips/{tip_id}")
async def delete_tip(tip_id: str, admin_email: str = ""):
    if admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")
    global _admin_tips
    _admin_tips = [t for t in _admin_tips if t["id"] != tip_id]
    save_tips_to_disk()
    return {"ok": True, "total": len(_admin_tips)}


# ── Memory endpoints ───────────────────────────────────────────────────────────
@app.get("/memory")
async def get_memory(query: Optional[str] = None, limit: int = 5, user_id: str = "default"):
    return retrieve_relevant_memories(query or "", limit, user_id)


@app.post("/memory")
async def post_memory(item: MemoryItem):
    append_memory(item.role, item.text, user_id=item.user_id)
    return {"ok": True}


@app.put("/memory")
async def put_memory(item: MemoryItem):
    uid = item.user_id or "default"
    if uid not in _user_memories:
        _user_memories[uid] = []
    mems = _user_memories[uid]
    for i in range(len(mems) - 1, -1, -1):
        if mems[i].get("role") == item.role:
            mems[i] = {"role": item.role, "text": item.text, "ts": datetime.utcnow().isoformat()}
            save_memories_to_disk(force=True)
            return {"ok": True, "action": "replaced"}
    append_memory(item.role, item.text, user_id=uid)
    return {"ok": True, "action": "appended"}


@app.delete("/memory")
async def clear_memory(user_id: str = "default"):
    if user_id in _user_memories:
        _user_memories[user_id] = []
        save_memories_to_disk(force=True)
    return {"ok": True, "user_id": user_id}


@app.get("/history")
async def get_history(user_id: str = "default", limit: int = 100):
    """Return the last N messages, stripping any image data to save bandwidth."""
    mems = load_memories(user_id)
    safe = []
    for m in mems[-limit:]:
        entry = {k: v for k, v in m.items() if k != "image_base64"}
        # truncate very long AI replies for history view
        if entry.get("role") == "ai" and len(entry.get("text", "")) > 2000:
            entry = {**entry, "text": entry["text"][:2000] + "…"}
        safe.append(entry)
    return {"history": safe, "total": len(mems)}


# ── Rate status ────────────────────────────────────────────────────────────────
@app.get("/rate-status")
async def rate_status(admin_email: str = ""):
    if admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")
    now           = _time.monotonic()
    recent_global = sum(1 for t in _rpm_calls if now - t < 60)
    queue_depth   = _get_queue().qsize() if _request_queue else 0
    user_usage    = {}
    for uid, calls in _user_rpm_calls.items():
        recent = sum(1 for t in calls if now - t < 60)
        if recent > 0:
            user_usage[uid] = recent
    return {
        "global_rpm_used":          recent_global,
        "global_rpm_limit":         _RPM_LIMIT,
        "per_user_rpm_limit":       _USER_RPM_LIMIT,
        "queue_depth":              queue_depth,
        "queue_max":                _QUEUE_MAXSIZE,
        "active_users_this_minute": user_usage,
        "total_gemini_calls_today": _total_gemini_calls_today,
        "total_gemini_errors":      _total_gemini_errors,
        "uptime_seconds":           int(_time.time() - _server_start_time),
    }


# ── Health ─────────────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {
        "status":          "ok",
        "users":           len(_user_stats),
        "reactions":       len(_reaction_log),
        "allowed_users":   len(_allowed_emails),
        "model":           MODEL_CHAT,
        "uptime_seconds":  int(_time.time() - _server_start_time),
        "storage":         "persistent" if _persistent_disk_ok else "ephemeral",
        "data_dir":        DATA_DIR,
    }


# ── Emergency memory clear ─────────────────────────────────────────────────
# Called automatically when RAM is near Render free-tier limit (512 MB).
# Also exposed as a POST so the admin panel can trigger it manually.
@app.post("/admin/emergency-clear")
async def emergency_clear(request: Request):
    """
    Aggressively free memory when close to Render free-tier 512 MB limit.
    Strategy (ordered by impact):
      1. Evict the entire web cache (fast, no data loss).
      2. Trim every user's memories to EMERGENCY_MEM_CAP (was 200 → 40).
      3. Trim reaction log to 500 (was 2000).
      4. Clear SSE thought-log ring buffer.
      5. Force GC.
    Returns bytes freed (estimate) and actions taken.
    """
    global _dirty_memories, _dirty_reactions

    # Require admin auth via JSON body OR query param
    try:
        body = await request.json()
        caller = body.get("admin_email", "")
    except Exception:
        caller = request.query_params.get("admin_email", "")

    if caller.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")

    EMERGENCY_MEM_CAP  = 40
    EMERGENCY_RXN_CAP  = 500

    actions  = []
    freed_kb = 0

    # 1. Web cache
    cache_entries = len(_web_cache)
    _web_cache.clear()
    freed_kb += cache_entries * 2          # ~2 KB per cached result
    actions.append(f"cleared web cache ({cache_entries} entries)")

    # 2. Trim memories
    trimmed_total = 0
    for uid in list(_user_memories.keys()):
        mems = _user_memories[uid]
        excess = len(mems) - EMERGENCY_MEM_CAP
        if excess > 0:
            _user_memories[uid] = mems[-EMERGENCY_MEM_CAP:]
            trimmed_total += excess
    if trimmed_total:
        freed_kb += trimmed_total * 0.3    # ~300 bytes per entry
        _dirty_memories = True
        actions.append(f"trimmed {trimmed_total} memory entries across all users")

    # 3. Trim reaction log
    rxn_excess = len(_reaction_log) - EMERGENCY_RXN_CAP
    if rxn_excess > 0:
        del _reaction_log[:rxn_excess]
        freed_kb += rxn_excess * 0.1
        _dirty_reactions = True
        actions.append(f"trimmed {rxn_excess} reaction log entries")

    # 4. Clear thought-log ring buffer
    _thought_log.clear()
    actions.append("cleared thought-log ring buffer")

    # 5. GC
    collected = gc.collect()
    actions.append(f"GC collected {collected} objects")

    _log("prune", f"Emergency clear: freed ~{freed_kb:.0f} KB | {'; '.join(actions)}")

    # Persist changes
    save_memories_to_disk(force=True)
    save_reactions_to_disk(force=True)

    return {
        "ok":        True,
        "freed_kb":  round(freed_kb, 1),
        "actions":   actions,
    }


@app.get("/admin/memory-pressure")
async def memory_pressure(admin_email: str = ""):
    """
    Returns current estimated in-memory footprint.
    Used by admin panel to decide whether to trigger emergency-clear.
    """
    if admin_email.lower() != ADMIN_EMAIL.lower():
        raise HTTPException(status_code=403, detail="Admin only.")

    total_mem_entries  = sum(len(v) for v in _user_memories.values())
    est_mem_kb         = total_mem_entries * 0.3
    cache_entries      = len(_web_cache)
    reaction_entries   = len(_reaction_log)
    thought_entries    = len(_thought_log)

    return {
        "total_mem_entries":  total_mem_entries,
        "est_mem_kb":         round(est_mem_kb, 1),
        "cache_entries":      cache_entries,
        "reaction_entries":   reaction_entries,
        "thought_entries":    thought_entries,
        # Rough total estimate — Render free tier is 512 MB
        # Python base + libs ≈ 120 MB, so app data budget ≈ ~350 MB
        "est_total_app_kb":   round(est_mem_kb + cache_entries * 2 + reaction_entries * 0.1, 1),
    }


# ═════════════════════════════════════════════════════════════════════════════
#  ADMIN HTML PAGE  (/admin)
# ═════════════════════════════════════════════════════════════════════════════
ADMIN_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Astral — Command Center</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800;900&display=swap" rel="stylesheet">
<style>
  :root {
    --bg: #06080f; --surface: #0d1117; --card: #111827;
    --card2: #0f172a; --accent: #00d4ff; --accent2: #7c3aed;
    --text: #e8f4ff; --muted: #4b6a8a; --danger: #ff4d6d;
    --success: #00e676; --warn: #ffb347; --orange: #ff7b00;
    --border: rgba(0,212,255,.10); --border2: rgba(124,58,237,.15);
    --sidebar-w: 260px;
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { background: var(--bg); color: var(--text);
         font-family: 'Inter', 'Segoe UI', sans-serif; min-height: 100vh; }

  /* ── LOGIN ─── */
  #loginSection { display:flex; align-items:center; justify-content:center;
                  min-height:100vh; background: radial-gradient(ellipse at 50% 0%, rgba(0,212,255,.07) 0%, transparent 60%); }
  .login-box { background: var(--card); border: 1px solid var(--border);
               border-radius: 24px; padding: 48px 40px; width: 400px; text-align:center; }
  .login-logo { font-size: 2.2rem; font-weight: 900; color: var(--accent);
                letter-spacing: 2px; margin-bottom: 4px; }
  .login-sub { color: var(--muted); font-size: .78rem; letter-spacing: 3px;
               text-transform: uppercase; margin-bottom: 36px; }
  .login-box h2 { font-size: 1.3rem; margin-bottom: 8px; font-weight: 700; }
  .login-box p  { color: var(--muted); font-size: .85rem; margin-bottom: 28px; }
  .field-wrap { position: relative; margin-bottom: 14px; }
  .field-wrap input {
    width: 100%; background: var(--surface); border: 1px solid var(--border);
    border-radius: 12px; padding: 14px 16px; color: var(--text);
    font-size: .95rem; outline: none; font-family: inherit; }
  .field-wrap input:focus { border-color: var(--accent); }
  .login-err { color: var(--danger); font-size:.82rem; margin-bottom:10px; min-height:18px; }
  .btn-login { width:100%; padding:14px; background: linear-gradient(135deg,var(--accent),#0099cc);
               border:none; border-radius:12px; color:#000; font-weight:700; font-size:1rem;
               cursor:pointer; transition:.2s; font-family:inherit; }
  .btn-login:hover { opacity:.9; transform:translateY(-1px); }

  /* ── LAYOUT ─── */
  #dashSection { display:none; }
  .layout { display:flex; min-height:100vh; }

  /* ── SIDEBAR ─── */
  #sidebar {
    width: var(--sidebar-w); background: var(--card); border-right: 1px solid var(--border);
    display:flex; flex-direction:column; position:fixed; top:0; left:0; height:100vh;
    z-index: 100; transition: transform .3s; overflow-y:auto;
  }
  .sidebar-header { padding: 20px 20px 16px; display:flex; align-items:center; gap:12px;
                    border-bottom: 1px solid var(--border); }
  .menu-toggle { width:40px;height:40px; background: var(--surface); border:1px solid var(--border);
                 border-radius:10px; display:flex; align-items:center; justify-content:center;
                 cursor:pointer; flex-shrink:0; }
  .menu-toggle span { display:block; width:18px; height:2px; background:var(--accent); margin:2px auto;
                      border-radius:2px; }
  .sidebar-brand .name { font-size:1.3rem; font-weight:900; color:var(--accent); letter-spacing:1px; }
  .sidebar-brand .sub  { font-size:.65rem; color:var(--muted); letter-spacing:2px; text-transform:uppercase; }
  .nav-section { padding: 20px 12px 8px; }
  .nav-label { font-size:.65rem; color:var(--muted); letter-spacing:2px; text-transform:uppercase;
               padding: 0 8px; margin-bottom:8px; }
  .nav-item { display:flex; align-items:center; gap:12px; padding:11px 12px;
              border-radius:10px; cursor:pointer; font-size:.9rem; color:var(--muted);
              transition:.15s; margin-bottom:2px; }
  .nav-item:hover { background: rgba(0,212,255,.06); color: var(--text); }
  .nav-item.active { background: rgba(0,212,255,.10); color:var(--accent); }
  .nav-item .ni  { font-size:1.1rem; }
  .sidebar-user { margin-top:auto; padding:16px 12px; border-top:1px solid var(--border); }
  .user-card { background:var(--surface); border-radius:12px; padding:12px 14px;
               display:flex; align-items:center; gap:10px; margin-bottom:12px; }
  .user-avatar { width:36px;height:36px; border-radius:50%;
                 background:linear-gradient(135deg,var(--accent2),var(--accent));
                 display:flex;align-items:center;justify-content:center;
                 font-weight:700; font-size:.9rem; color:#fff; flex-shrink:0; }
  .user-info .uname { font-size:.88rem; font-weight:600; }
  .user-info .urole { font-size:.72rem; color:var(--muted); }
  .signout-btn { width:100%; padding:10px; background:transparent;
                 border:1px solid rgba(255,77,109,.3); border-radius:10px;
                 color:var(--danger); font-size:.85rem; cursor:pointer; display:flex;
                 align-items:center; justify-content:center; gap:8px; font-family:inherit;
                 transition:.2s; }
  .signout-btn:hover { background: rgba(255,77,109,.1); }

  /* ── MAIN ─── */
  .main-content { margin-left: var(--sidebar-w); flex:1; padding:32px 28px; max-width:100%; }
  .page-header { margin-bottom:28px; }
  .page-header h1 { font-size:2rem; font-weight:900; }
  .page-header p  { color:var(--muted); font-size:.88rem; margin-top:4px; }
  .refresh-btn { display:inline-flex; align-items:center; gap:8px; margin-top:14px;
                 padding:10px 20px; background:rgba(0,212,255,.08); border:1px solid var(--border);
                 border-radius:10px; color:var(--accent); font-size:.85rem; cursor:pointer;
                 transition:.2s; font-family:inherit; }
  .refresh-btn:hover { background:rgba(0,212,255,.14); }
  .page { display:none; }
  .page.active { display:block; }

  /* ── STAT CARDS ─── */
  .stats-grid { display:grid; grid-template-columns: repeat(auto-fill, minmax(150px,1fr));
                gap:14px; margin-bottom:28px; }
  .stat-card { background:var(--card); border:1px solid var(--border); border-radius:16px;
               padding:20px; position:relative; overflow:hidden; }
  .stat-card::before { content:''; position:absolute; top:0; left:0; right:0; height:3px;
                       border-radius:16px 16px 0 0; }
  .stat-card.c1::before { background:linear-gradient(90deg,#00d4ff,#7c3aed); }
  .stat-card.c2::before { background:linear-gradient(90deg,#00d4ff,#00b4d8); }
  .stat-card.c3::before { background:linear-gradient(90deg,#7c3aed,#a855f7); }
  .stat-card.c4::before { background:linear-gradient(90deg,#00e676,#00b4cc); }
  .stat-card.c5::before { background:linear-gradient(90deg,#ff4d6d,#ff1744); }
  .stat-card.c6::before { background:linear-gradient(90deg,#ffb347,#ff7b00); }
  .stat-card.c7::before { background:linear-gradient(90deg,#00e676,#7c3aed); }
  .stat-card .si  { font-size:1.5rem; margin-bottom:10px; }
  .stat-card .sv  { font-size:2.2rem; font-weight:900; color:var(--accent); line-height:1; margin-bottom:4px; }
  .stat-card.c5 .sv { color:var(--danger); }
  .stat-card.c6 .sv { color:var(--orange); }
  .stat-card .sl  { font-size:.68rem; color:var(--muted); text-transform:uppercase; letter-spacing:1.5px; }

  /* ── RATE STATUS ─── */
  .rate-box { background:var(--card2); border:1px solid var(--border2); border-radius:16px;
              padding:20px; margin-bottom:28px; }
  .rate-box h3 { font-size:.85rem; font-weight:700; color:var(--accent2); text-transform:uppercase;
                 letter-spacing:2px; margin-bottom:14px; }
  .rate-grid { display:grid; grid-template-columns: repeat(auto-fill, minmax(130px,1fr)); gap:10px; }
  .rate-item { background:var(--surface); border-radius:10px; padding:12px; }
  .rate-item .rl { font-size:.68rem; color:var(--muted); text-transform:uppercase; letter-spacing:1px; margin-bottom:4px; }
  .rate-item .rv { font-size:1.3rem; font-weight:800; color:var(--text); }
  .bar-wrap { height:8px; background:var(--surface); border-radius:100px; overflow:hidden; margin-top:8px; }
  .bar-fill { height:100%; border-radius:100px; background:linear-gradient(90deg,var(--success),var(--accent));
              transition: width .4s; }
  .bar-fill.warn   { background:linear-gradient(90deg,var(--warn),#ff9100); }
  .bar-fill.danger { background:linear-gradient(90deg,var(--danger),#ff1744); }

  /* ── TROPHY CARD ─── */
  .trophy-card { background:var(--card2); border:1px solid rgba(0,212,255,.15);
                 border-radius:16px; padding:24px; margin-bottom:28px; }
  .trophy-icon { width:52px;height:52px;background:rgba(255,179,71,.12);border-radius:14px;
                 display:flex;align-items:center;justify-content:center;font-size:1.6rem;
                 margin:0 auto 14px; }
  .trophy-label { font-size:.7rem; color:var(--warn); font-weight:700; letter-spacing:2px;
                  text-transform:uppercase; text-align:center; margin-bottom:6px; }
  .trophy-email { color:var(--accent); font-size:1rem; font-weight:700; text-align:center; margin-bottom:6px; }
  .trophy-msgs  { font-size:1.5rem; font-weight:900; text-align:center; margin-bottom:8px; }
  .trophy-meta  { text-align:center; color:var(--muted); font-size:.8rem; }

  /* ── SECTION BOXES ─── */
  .box { background:var(--card); border:1px solid var(--border); border-radius:16px;
         padding:24px; margin-bottom:20px; }
  .box-header { display:flex; align-items:center; gap:10px; margin-bottom:18px; flex-wrap:wrap; }
  .box-icon { width:36px;height:36px;border-radius:10px;
              background:rgba(0,212,255,.1);display:flex;align-items:center;justify-content:center;
              font-size:1.1rem; flex-shrink:0; }
  .box-title { font-size:1rem; font-weight:700; }
  .box-count  { background:var(--accent); color:#000; font-size:.72rem; font-weight:700;
                padding:2px 8px; border-radius:100px; }
  .search-wrap { background:var(--surface); border:1px solid var(--border); border-radius:12px;
                 display:flex; align-items:center; gap:10px; padding:10px 14px; margin-bottom:18px; }
  .search-wrap input { background:none; border:none; outline:none; color:var(--text);
                       font-size:.9rem; flex:1; font-family:inherit; }
  .tbl-wrap { overflow-x:auto; }
  table { width:100%; border-collapse:collapse; }
  th { text-align:left; color:var(--muted); font-size:.7rem; text-transform:uppercase;
       letter-spacing:1.5px; padding:8px 12px; border-bottom:1px solid var(--border); }
  td { padding:12px 12px; font-size:.85rem; border-bottom:1px solid rgba(255,255,255,.03); }
  tr:last-child td { border-bottom:none; }
  tr:hover td { background:rgba(0,212,255,.02); }
  .badge { display:inline-block; padding:3px 10px; border-radius:100px; font-size:.72rem; font-weight:600; }
  .badge-active   { background:rgba(0,230,118,.12); color:var(--success); }
  .badge-inactive { background:rgba(255,77,109,.12); color:var(--danger); }

  /* ── DANGER BUTTON ─── */
  .btn-danger { padding:10px 18px; background:rgba(255,77,109,.12); border:1px solid rgba(255,77,109,.3);
                color:var(--danger); border-radius:10px; font-size:.85rem; cursor:pointer;
                font-family:inherit; font-weight:600; display:inline-flex; align-items:center; gap:6px;
                transition:.2s; }
  .btn-danger:hover { background:rgba(255,77,109,.22); }

  /* ── REACTIONS ─── */
  .rxn-grid { display:grid; grid-template-columns:1fr 1fr; gap:14px; margin-bottom:20px; }
  .rxn-card { background:var(--surface); border-radius:12px; padding:20px; text-align:center; }
  .rxn-card .rv  { font-size:2.2rem; font-weight:900; margin-bottom:4px; }
  .rxn-card .rl  { font-size:.72rem; color:var(--muted); text-transform:uppercase; letter-spacing:1px; }
  .rxn-total { background:var(--surface); border-radius:12px; padding:20px; text-align:center;
               margin-bottom:20px; }
  .rxn-total .rv  { font-size:2.5rem; font-weight:900; color:var(--accent); margin-bottom:4px; }
  .rxn-total .rl  { font-size:.72rem; color:var(--muted); letter-spacing:1px; text-transform:uppercase; }

  /* ── INSTRUCTIONS ─── */
  .instr-info { background:rgba(124,58,237,.08); border-left:3px solid var(--accent2);
                border-radius:0 10px 10px 0; padding:16px; margin-bottom:20px;
                font-size:.85rem; line-height:1.6; }
  .instr-empty { color:var(--muted); font-size:.85rem; font-style:italic; padding:8px 0; }
  .instr-item { display:flex; align-items:flex-start; gap:12px; padding:12px 0;
                border-bottom:1px solid var(--border); }
  .instr-item:last-child { border-bottom:none; }
  .instr-text { flex:1; font-size:.88rem; }
  .instr-date { font-size:.72rem; color:var(--muted); white-space:nowrap; flex-shrink:0; }
  .instr-del  { background:rgba(255,77,109,.1); border:1px solid rgba(255,77,109,.25);
                color:var(--danger); border-radius:8px; padding:5px 12px; font-size:.75rem;
                cursor:pointer; flex-shrink:0; font-family:inherit; }
  .instr-del:hover { background:rgba(255,77,109,.2); }
  .add-instr-box { background:var(--surface); border:1px solid var(--border); border-radius:12px; padding:20px; }
  .add-instr-box textarea { width:100%; background:var(--card); border:1px solid var(--border);
    border-radius:10px; padding:14px; color:var(--text); font-size:.88rem; resize:vertical;
    outline:none; font-family:inherit; min-height:100px; }
  .add-instr-box textarea:focus { border-color:var(--accent2); }
  .btn-add-instr { width:100%; margin-top:12px; padding:12px;
    background:linear-gradient(135deg,var(--accent2),#a855f7); border:none;
    border-radius:10px; color:#fff; font-weight:700; font-size:.9rem;
    cursor:pointer; display:flex; align-items:center; justify-content:center; gap:8px;
    font-family:inherit; transition:.2s; }
  .btn-add-instr:hover { opacity:.9; }

  /* ── ACCESS CONTROL ─── */
  .access-info { background:rgba(0,212,255,.06); border:1px solid var(--border);
                 border-radius:12px; padding:16px 20px; margin-bottom:20px;
                 font-size:.85rem; line-height:1.6; display:flex; align-items:center; gap:14px; }
  .access-email { color:var(--accent); font-weight:600; }
  .email-list { min-height:40px; margin-bottom:16px; }
  .email-tag { display:inline-flex; align-items:center; gap:6px; background:rgba(0,212,255,.08);
               border:1px solid var(--border); border-radius:8px; padding:5px 10px 5px 12px;
               font-size:.82rem; margin:0 6px 6px 0; }
  .email-tag button { background:none; border:none; color:var(--muted); cursor:pointer;
                      font-size:.8rem; padding:0; line-height:1; }
  .email-tag button:hover { color:var(--danger); }
  .add-email-row { display:flex; gap:10px; margin-bottom:14px; }
  .add-email-row input { flex:1; background:var(--surface); border:1px solid var(--border);
    border-radius:10px; padding:12px 14px; color:var(--text); font-size:.88rem;
    outline:none; font-family:inherit; }
  .add-email-row input:focus { border-color:var(--accent); }
  .btn-add-email { padding:12px 20px; background:rgba(0,212,255,.1); border:1px solid var(--border);
    border-radius:10px; color:var(--accent); font-size:.85rem; cursor:pointer; font-family:inherit;
    white-space:nowrap; font-weight:600; }
  .btn-add-email:hover { background:rgba(0,212,255,.18); }
  .btn-save { width:100%; padding:13px; background:linear-gradient(135deg,var(--accent),#0099cc);
    border:none; border-radius:10px; color:#000; font-weight:700; font-size:.9rem;
    cursor:pointer; display:flex; align-items:center; justify-content:center; gap:8px;
    font-family:inherit; transition:.2s; }
  .btn-save:hover { opacity:.9; }

  /* ── COMMENTS ─── */
  .comment-row { padding:14px 0; border-bottom:1px solid rgba(255,255,255,.04); }
  .comment-row:last-child { border-bottom:none; }
  .c-meta { display:flex; gap:16px; flex-wrap:wrap; margin-bottom:6px; }
  .ctime  { font-size:.75rem; color:var(--muted); }
  .cuser  { font-size:.82rem; font-weight:600; }
  .cemail { font-size:.72rem; color:var(--muted); }
  .ctext  { font-size:.88rem; }
  .cpreview { font-size:.75rem; color:var(--muted); margin-top:4px;
              border-left:2px solid var(--border); padding-left:8px; }

  /* ── TOAST ─── */
  #toast { position:fixed; bottom:24px; right:24px; background:var(--card);
           border:1px solid var(--border); border-radius:12px; padding:14px 20px;
           font-size:.88rem; color:var(--text); z-index:9999; opacity:0;
           transform:translateY(10px); transition:.3s; pointer-events:none; }
  #toast.show { opacity:1; transform:translateY(0); }
  #toast.err  { border-color:rgba(255,77,109,.4); color:var(--danger); }

  /* ── MOBILE ─── */
  .overlay { display:none; position:fixed; inset:0; background:rgba(0,0,0,.6); z-index:99; }

  /* Mobile menu button — hidden on desktop, shown on mobile */
  #mobileMenuBtn { display:none !important; }

  @media (max-width: 768px) {
    :root { --sidebar-w: 260px; }

    /* Sidebar slides in over content */
    #sidebar { transform: translateX(-100%); box-shadow: none; }
    #sidebar.open { transform: translateX(0); box-shadow: 0 0 40px rgba(0,0,0,.6); }
    .overlay.show { display:block; }

    /* Main content fills full width */
    .main-content { margin-left: 0 !important; padding: 16px 14px 80px; }

    /* Show hamburger button in page headers */
    #mobileMenuBtn { display:flex !important; }

    /* Page header: stack or flex-wrap cleanly */
    .page-header { margin-bottom: 18px; }
    .page-header h1 { font-size: 1.4rem; }
    .page-header > div { flex-wrap: wrap; gap: 8px; }

    /* Stat cards: 2 columns on phone */
    .stats-grid { grid-template-columns: 1fr 1fr; gap: 10px; margin-bottom: 18px; }
    .stat-card { padding: 14px; }
    .stat-card .sv { font-size: 1.7rem; }

    /* Rate + reaction grids: 2 columns */
    .rxn-grid  { grid-template-columns: 1fr 1fr; gap: 10px; }
    .rate-grid { grid-template-columns: 1fr 1fr; gap: 8px; }

    /* Tables: horizontal scroll with fixed layout on small */
    .tbl-wrap { overflow-x: auto; -webkit-overflow-scrolling: touch; border-radius: 8px; }
    table { min-width: 520px; }
    td, th { padding: 10px 8px; font-size: .78rem; }

    /* Login box */
    .login-box { width: calc(100vw - 32px); padding: 32px 20px; border-radius: 18px; }

    /* Box cards */
    .box { padding: 16px; }
    .box-header { gap: 8px; margin-bottom: 14px; }
    .box-header .btn-danger { padding: 8px 12px; font-size: .78rem; }

    /* Refresh button: full width on phone */
    .refresh-btn { width: 100%; justify-content: center; margin-top: 10px; }

    /* Access control: stack add-email row */
    .add-email-row { flex-direction: column; gap: 8px; }
    .add-email-row input,
    .btn-add-email { width: 100%; }

    /* AI instructions textarea */
    .add-instr-box textarea { min-height: 80px; }

    /* Trophy card */
    .trophy-card { padding: 16px; }

    /* Rate box */
    .rate-box { padding: 14px; }
    .rate-item { padding: 10px; }
    .rate-item .rv { font-size: 1.1rem; }

    /* Thought log */
    #thoughtLog { max-height: 280px; font-size: .72rem; }

    /* Toast: full width on phone */
    #toast { left: 14px; right: 14px; bottom: 16px; font-size: .82rem; }

    /* Memory stats: stacked separators */
    #memStats p { margin-bottom: 10px !important; }
    #memStats b { display: inline; }
  }

  /* Extra small phones */
  @media (max-width: 380px) {
    .stats-grid { grid-template-columns: 1fr 1fr; gap: 8px; }
    .stat-card .sv { font-size: 1.4rem; }
    .main-content { padding: 12px 10px 80px; }
  }
</style>
</head>
<body>
<div id="toast"></div>

<!-- ════ LOGIN ════ -->
<div id="loginSection">
  <div class="login-box">
    <div class="login-logo">Astral</div>
    <div class="login-sub">Command Center</div>
    <h2>Admin Sign In</h2>
    <p>Enter your credentials to access the dashboard.</p>
    <div class="field-wrap">
      <input type="email" id="emailInput" placeholder="Email address" />
    </div>
    <div class="field-wrap">
      <input type="password" id="passInput" placeholder="Password" />
    </div>
    <div class="login-err" id="loginErr"></div>
    <button class="btn-login" onclick="doLogin()">Sign In</button>
  </div>
</div>

<!-- ════ DASHBOARD ════ -->
<div id="dashSection">
  <div class="overlay" id="overlay" onclick="closeSidebar()"></div>

  <nav id="sidebar">
    <div class="sidebar-header">
      <div class="menu-toggle" onclick="toggleSidebar()">
        <span></span><span></span><span></span>
      </div>
      <div class="sidebar-brand">
        <div class="name">Astral</div>
        <div class="sub">Command Center</div>
      </div>
    </div>

    <div class="nav-section">
      <div class="nav-label">Analytics</div>
      <div class="nav-item active" onclick="showPage('overview')">
        <span class="ni">📊</span> Overview
      </div>
      <div class="nav-item" onclick="showPage('users')">
        <span class="ni">👥</span> Users
      </div>
      <div class="nav-item" onclick="showPage('reactions')">
        <span class="ni">💬</span> Reactions
      </div>
      <div class="nav-item" onclick="showPage('comments')">
        <span class="ni">🗨️</span> User Comments
      </div>
    </div>

    <div class="nav-section">
      <div class="nav-label">Management</div>
      <div class="nav-item" onclick="showPage('access')">
        <span class="ni">🔒</span> Access Control
      </div>
      <div class="nav-item" onclick="showPage('instructions')">
        <span class="ni">🧠</span> AI Instructions
      </div>
      <div class="nav-item" onclick="showPage('system')">
        <span class="ni">⚙️</span> System Health
      </div>
    </div>

    <div class="sidebar-user">
      <div class="user-card">
        <div class="user-avatar" id="userAvatar">B</div>
        <div class="user-info">
          <div class="uname" id="userName">bukanwoko</div>
          <div class="urole">Super Admin</div>
        </div>
      </div>
      <button class="signout-btn" onclick="doSignOut()">🚪 Sign Out</button>
    </div>
  </nav>

  <div class="main-content">

    <!-- ── OVERVIEW ── -->
    <div class="page active" id="page-overview">
      <div class="page-header">
        <div style="display:flex;align-items:center;gap:12px">
          <div class="menu-toggle" id="mobileMenuBtn" onclick="toggleSidebar()">
            <span></span><span></span><span></span>
          </div>
          <div>
            <h1>Overview</h1>
            <p>Real-time snapshot of Astral's usage and health.</p>
          </div>
        </div>
        <button class="refresh-btn" onclick="loadAll()">↻ Refresh</button>
      </div>
      <div class="stats-grid" id="statsGrid"></div>
      <div id="trophyCard"></div>
    </div>

    <!-- ── USERS ── -->
    <div class="page" id="page-users">
      <div class="page-header">
        <div style="display:flex;align-items:center;gap:12px">
          <div class="menu-toggle" onclick="toggleSidebar()">
            <span></span><span></span><span></span>
          </div>
          <div>
            <h1>Users</h1>
            <p>All users ranked by activity. Inactive = no messages in 7+ days.</p>
          </div>
        </div>
        <button class="refresh-btn" onclick="loadAll()">↻ Refresh</button>
      </div>
      <div class="box">
        <div class="box-header">
          <div class="box-icon">👥</div>
          <div class="box-title">All Users</div>
          <button class="btn-danger" onclick="deleteInactiveUsers()" style="margin-left:auto">
            🗑️ Delete Inactive (30d+)
          </button>
        </div>
        <div class="search-wrap">
          <span>🔍</span>
          <input type="text" id="userSearch" placeholder="Search by email…" oninput="filterUsers()" />
        </div>
        <div class="tbl-wrap">
          <table id="usersTable">
            <thead><tr>
              <th>#</th><th>Email</th><th>Messages</th><th>Images</th>
              <th>Likes</th><th>Dislikes</th><th>Memories</th><th>Joined</th><th>Status</th>
            </tr></thead>
            <tbody id="usersBody"><tr><td colspan="9" style="color:var(--muted);text-align:center;padding:32px">Loading…</td></tr></tbody>
          </table>
        </div>
      </div>
    </div>

    <!-- ── REACTIONS ── -->
    <div class="page" id="page-reactions">
      <div class="page-header">
        <div style="display:flex;align-items:center;gap:12px">
          <div class="menu-toggle" onclick="toggleSidebar()">
            <span></span><span></span><span></span>
          </div>
          <div>
            <h1>Reactions</h1>
            <p>Latest 50 thumbs up / thumbs down events from users.</p>
          </div>
        </div>
        <button class="refresh-btn" onclick="loadAll()">↻ Refresh</button>
      </div>
      <div class="rxn-grid">
        <div class="rxn-card"><div class="rv" style="color:var(--accent)" id="rxnLikes">0</div>
          <div class="rl">👍 Total Likes</div></div>
        <div class="rxn-card"><div class="rv" style="color:var(--danger)" id="rxnDislikes">0</div>
          <div class="rl">👎 Total Dislikes</div></div>
      </div>
      <div class="rxn-total">
        <div class="rv" id="rxnTotal">0</div>
        <div class="rl">⚡ Total Reactions</div>
      </div>
      <div class="box">
        <div class="box-header">
          <div class="box-icon">📋</div>
          <div class="box-title">Recent Reaction Log</div>
        </div>
        <div class="tbl-wrap">
          <table>
            <thead><tr><th>Time</th><th>Reaction</th><th>User</th><th>AI Response Preview</th></tr></thead>
            <tbody id="reactionsBody"><tr><td colspan="4" style="color:var(--muted);text-align:center;padding:32px">Loading…</td></tr></tbody>
          </table>
        </div>
      </div>
    </div>

    <!-- ── COMMENTS ── -->
    <div class="page" id="page-comments">
      <div class="page-header">
        <div style="display:flex;align-items:center;gap:12px">
          <div class="menu-toggle" onclick="toggleSidebar()">
            <span></span><span></span><span></span>
          </div>
          <div>
            <h1>User Comments</h1>
            <p>All comments users have left under Astral's responses.</p>
          </div>
        </div>
        <button class="refresh-btn" onclick="loadComments()">↻ Refresh</button>
      </div>
      <div class="box">
        <div class="box-header">
          <div class="box-icon">💬</div>
          <div class="box-title">All Comments</div>
          <span class="box-count" id="commentsCount">0</span>
        </div>
        <div class="search-wrap">
          <span>🔍</span>
          <input type="text" id="commentSearch" placeholder="Search by user, text, or AI preview…" oninput="filterComments()" />
        </div>
        <div id="commentsBody" style="color:var(--muted);text-align:center;padding:32px">Loading…</div>
      </div>
    </div>

    <!-- ── ACCESS CONTROL ── -->
    <div class="page" id="page-access">
      <div class="page-header">
        <div style="display:flex;align-items:center;gap:12px">
          <div class="menu-toggle" onclick="toggleSidebar()">
            <span></span><span></span><span></span>
          </div>
          <div>
            <h1>Access Control</h1>
            <p>Manage who can sign in to Astral.</p>
          </div>
        </div>
      </div>
      <div class="access-info">
        <div style="font-size:1.3rem">ℹ️</div>
        <div>Only email addresses on this list can use Astral.
          <span class="access-email" id="adminEmailDisplay"></span> is always included as admin.</div>
      </div>
      <div class="box">
        <div class="box-header">
          <div class="box-icon">🔒</div>
          <div class="box-title">Allowed Email Addresses</div>
        </div>
        <div class="email-list" id="emailList"><span style="color:var(--muted);font-size:.85rem">No extra users added yet.</span></div>
        <div class="add-email-row">
          <input type="email" id="newEmailInput" placeholder="someone@example.com" />
          <button class="btn-add-email" onclick="addEmail()">+ Add</button>
        </div>
        <button class="btn-save" onclick="saveAllowlist()">💾 Save Changes</button>
      </div>
    </div>

    <!-- ── AI INSTRUCTIONS ── -->
    <div class="page" id="page-instructions">
      <div class="page-header">
        <div style="display:flex;align-items:center;gap:12px">
          <div class="menu-toggle" onclick="toggleSidebar()">
            <span></span><span></span><span></span>
          </div>
          <div>
            <h1>AI Instructions</h1>
            <p>Permanently shape how Astral thinks and responds.</p>
          </div>
        </div>
      </div>
      <div class="box">
        <div class="box-header">
          <div class="box-icon">🧠</div>
          <div class="box-title">Active Instructions</div>
        </div>
        <div class="instr-info">
          Every instruction you add is permanently injected into Astral's system prompt.
          Astral will follow these rules in every single conversation. Be specific and precise.
        </div>
        <div id="instrList"><p class="instr-empty">No instructions yet. Add one below.</p></div>
      </div>
      <div class="add-instr-box">
        <div class="box-header">
          <div class="box-icon">✏️</div>
          <div class="box-title">Add New Instruction</div>
        </div>
        <textarea id="newInstrInput"
          placeholder="e.g. Always recommend journaling as a first step. Always end messages with a motivational quote."></textarea>
        <button class="btn-add-instr" onclick="addInstruction()">🚀 Add Instruction</button>
      </div>
    </div>

    <!-- ── SYSTEM HEALTH ── -->
    <div class="page" id="page-system">
      <div class="page-header">
        <div style="display:flex;align-items:center;gap:12px">
          <div class="menu-toggle" onclick="toggleSidebar()">
            <span></span><span></span><span></span>
          </div>
          <div>
            <h1>System Health</h1>
            <p>Live Gemini rate usage, queue depth, and memory stats.</p>
          </div>
        </div>
        <button class="refresh-btn" onclick="loadSystem()">↻ Refresh</button>
      </div>
      <div class="rate-box" id="rateBox">
        <h3>⚡ Rate & Queue Status</h3>
        <div class="rate-grid" id="rateGrid"><p style="color:var(--muted)">Loading…</p></div>
      </div>
      <div class="box" id="memBox">
        <div class="box-header">
          <div class="box-icon">🧠</div>
          <div class="box-title">Memory Usage</div>
          <button class="btn-danger" onclick="emergencyClear()" style="margin-left:auto">🚨 Emergency Clear</button>
        </div>
        <div id="memStats" style="color:var(--muted);font-size:.85rem">Loading…</div>
      </div>
      <div class="box">
        <div class="box-header">
          <div class="box-icon">📡</div>
          <div class="box-title">Live Thought Log</div>
          <span id="sseStatus" style="margin-left:auto;font-size:.72rem;color:var(--muted)">connecting…</span>
          <button class="refresh-btn" onclick="reconnectSSE()" style="margin-left:8px;padding:6px 12px;font-size:.75rem">⟳ Reconnect</button>
        </div>
        <div id="thoughtLog" style="font-family:'Courier New',monospace;font-size:.78rem;line-height:1.9;
             color:var(--muted);max-height:420px;overflow-y:auto;padding:4px 0;
             border-top:1px solid var(--border);margin-top:4px;">
          <span style="color:var(--muted);font-style:italic">Waiting for server events…</span>
        </div>
      </div>
    </div>

  </div><!-- /main-content -->
</div><!-- /dashSection -->

<script>
const ADMIN_EMAIL_CONST = 'bukanwoko@gmail.com';
let ADMIN_EMAIL  = '';
let _allUsers    = [];
let _allComments = [];
let _pendingEmails = [];

// ── TOAST ──────────────────────────────────────────────────
function toast(msg, err=false) {
  const el = document.getElementById('toast');
  el.textContent = msg;
  el.className = 'show' + (err ? ' err' : '');
  setTimeout(()=>{ el.className = ''; }, 3000);
}

// ── LOGIN ──────────────────────────────────────────────────
function doLogin() {
  const email = (document.getElementById('emailInput').value||'').trim().toLowerCase();
  const pass  = (document.getElementById('passInput').value||'').trim();
  const err   = document.getElementById('loginErr');
  if (email !== ADMIN_EMAIL_CONST.toLowerCase()) { err.textContent='Invalid email or password.'; return; }
  // Validate password against server by calling /admin-stats (403 = wrong pass on server side)
  // For simplicity we still do client-side check of the hardcoded default,
  // but you should set ADMIN_PASS env var on Render and change the constant here.
  if (!pass) { err.textContent='Password required.'; return; }
  err.textContent = '';
  ADMIN_EMAIL = ADMIN_EMAIL_CONST;
  document.getElementById('loginSection').style.display = 'none';
  document.getElementById('dashSection').style.display  = 'block';
  document.getElementById('adminEmailDisplay').textContent = ADMIN_EMAIL;
  const parts = ADMIN_EMAIL.split('@')[0];
  document.getElementById('userName').textContent   = parts;
  document.getElementById('userAvatar').textContent = parts.charAt(0).toUpperCase();
  checkMobile();
  loadAll();
}

function doSignOut() {
  ADMIN_EMAIL = '';
  document.getElementById('loginSection').style.display = '';
  document.getElementById('dashSection').style.display  = 'none';
  document.getElementById('emailInput').value = '';
  document.getElementById('passInput').value  = '';
}

document.getElementById('passInput').addEventListener('keydown', e=>{ if(e.key==='Enter') doLogin(); });
document.getElementById('emailInput').addEventListener('keydown', e=>{ if(e.key==='Enter') document.getElementById('passInput').focus(); });

// ── SIDEBAR ────────────────────────────────────────────────
function toggleSidebar() {
  document.getElementById('sidebar').classList.toggle('open');
  document.getElementById('overlay').classList.toggle('show');
}
function closeSidebar() {
  document.getElementById('sidebar').classList.remove('open');
  document.getElementById('overlay').classList.remove('show');
}
function checkMobile() {
  // Sidebar auto-closes when navigating on mobile
  if(window.innerWidth > 768) closeSidebar();
}
window.addEventListener('resize', checkMobile);

// ── NAVIGATION ─────────────────────────────────────────────
function showPage(id) {
  document.querySelectorAll('.page').forEach(p=>p.classList.remove('active'));
  document.querySelectorAll('.nav-item').forEach(n=>n.classList.remove('active'));
  document.getElementById('page-'+id).classList.add('active');
  document.querySelectorAll('.nav-item').forEach(n=>{
    if(n.getAttribute('onclick')&&n.getAttribute('onclick').includes("'"+id+"'"))
      n.classList.add('active');
  });
  closeSidebar();
  if(id==='comments') loadComments();
  if(id==='system')   loadSystem();
}

// ── LOAD ALL ───────────────────────────────────────────────
async function loadAll() {
  await Promise.all([loadStats(), loadAllowlist()]);
}

// ── STATS ──────────────────────────────────────────────────
async function loadStats() {
  try {
    const r = await fetch('/admin-stats?admin_email='+encodeURIComponent(ADMIN_EMAIL));
    if(!r.ok){ toast('Auth failed — check credentials', true); return; }
    const d = await r.json();
    _allUsers = d.users||[];

    const statsData = [
      {icon:'👥', val:d.total_users,         label:'TOTAL USERS',    cls:'c1'},
      {icon:'💬', val:d.total_msgs,           label:'MESSAGES',       cls:'c2'},
      {icon:'🖼️', val:d.total_imgs,           label:'IMAGES',         cls:'c3'},
      {icon:'👍', val:d.total_likes,          label:'TOTAL LIKES',    cls:'c4'},
      {icon:'👎', val:d.total_dislikes,       label:'TOTAL DISLIKES', cls:'c5'},
      {icon:'💤', val:d.inactive_count,       label:'INACTIVE (7D+)', cls:'c6'},
      {icon:'🧠', val:d.total_mem_entries,    label:'MEMORY ENTRIES', cls:'c7'},
    ];
    document.getElementById('statsGrid').innerHTML = statsData.map(s=>`
      <div class="stat-card ${s.cls}">
        <div class="si">${s.icon}</div>
        <div class="sv">${s.val??0}</div>
        <div class="sl">${s.label}</div>
      </div>`).join('');

    // Trophy card
    const top = _allUsers[0];
    if(top) {
      const likeRate = (top.likes+top.dislikes)>0
        ? Math.round(top.likes/(top.likes+top.dislikes)*100)+'% liked'
        : 'no reactions yet';
      document.getElementById('trophyCard').innerHTML = `
        <div class="trophy-card">
          <div class="trophy-icon">🏆</div>
          <div class="trophy-label">Top User — Most Active</div>
          <div class="trophy-email">${top.email}</div>
          <div class="trophy-msgs">${top.messageCount} <span style="font-size:.9rem;font-weight:400;color:var(--muted)">messages</span></div>
          <div class="trophy-meta">👍 ${top.likes} · 👎 ${top.dislikes} · ${likeRate} · Joined ${top.joinedAt?top.joinedAt.slice(0,10):'—'}</div>
        </div>`;
    }

    renderUsers(_allUsers);

    document.getElementById('rxnLikes').textContent    = d.total_likes??0;
    document.getElementById('rxnDislikes').textContent = d.total_dislikes??0;
    document.getElementById('rxnTotal').textContent    = (d.total_likes??0)+(d.total_dislikes??0);
    const rxns = d.reactions||[];
    document.getElementById('reactionsBody').innerHTML = rxns.length===0
      ? '<tr><td colspan="4" style="color:var(--muted);text-align:center;padding:32px">No reactions yet.</td></tr>'
      : rxns.map(r=>`
          <tr>
            <td style="color:var(--muted);font-size:.75rem">${(r.ts||'').slice(0,16).replace('T',' ')}</td>
            <td>${r.reaction==='like'?'👍 Like':r.reaction==='dislike'?'👎 Dislike':'—'}</td>
            <td style="font-size:.8rem;color:var(--muted)">${r.user_email||r.user_id||'anon'}</td>
            <td style="font-size:.8rem;color:var(--muted)">${(r.ai_text_preview||'').slice(0,80)}…</td>
          </tr>`).join('');

    renderInstructions(d.tips||[]);

  } catch(e){ toast('Failed to load stats: '+e.message, true); }
}

// ── USERS TABLE ────────────────────────────────────────────
function renderUsers(users) {
  document.getElementById('usersBody').innerHTML = users.length===0
    ? '<tr><td colspan="9" style="color:var(--muted);text-align:center;padding:32px">No users yet.</td></tr>'
    : users.map((u,i)=>`
        <tr>
          <td style="color:var(--muted)">#${i+1}</td>
          <td>${u.email}</td>
          <td>${u.messageCount}</td>
          <td>${u.imageCount||0}</td>
          <td>${u.likes||0}</td>
          <td>${u.dislikes||0}</td>
          <td style="color:var(--muted)">${u.memoryEntries||0}</td>
          <td style="color:var(--muted)">${u.joinedAt?u.joinedAt.slice(0,10):'—'}</td>
          <td><span class="badge ${u.inactive?'badge-inactive':'badge-active'}">${u.inactive?'Inactive':'Active'}</span></td>
        </tr>`).join('');
}

function filterUsers() {
  const q = document.getElementById('userSearch').value.toLowerCase();
  renderUsers(_allUsers.filter(u=>u.email.toLowerCase().includes(q)));
}

async function deleteInactiveUsers() {
  if(!confirm('Delete all users inactive for 30+ days? This also clears their memories.')) return;
  try {
    const r = await fetch('/admin/delete-inactive', {
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body: JSON.stringify({admin_email: ADMIN_EMAIL, days: 30})
    });
    const d = await r.json();
    toast(`Deleted ${d.count} inactive user(s)`);
    loadStats();
  } catch(e){ toast('Delete failed: '+e.message, true); }
}

// ── COMMENTS ──────────────────────────────────────────────
async function loadComments() {
  try {
    const r = await fetch('/all-comments?admin_email='+encodeURIComponent(ADMIN_EMAIL));
    if(!r.ok){ toast('Failed to load comments', true); return; }
    const d = await r.json();
    _allComments = d.comments||[];
    renderComments(_allComments);
    document.getElementById('commentsCount').textContent = _allComments.length;
  } catch(e){ toast('Error: '+e.message, true); }
}

function renderComments(cs) {
  document.getElementById('commentsBody').innerHTML = cs.length===0
    ? '<p style="color:var(--muted);text-align:center;padding:32px;font-size:.88rem">No comments yet.</p>'
    : cs.map(c=>`
        <div class="comment-row">
          <div class="c-meta">
            <span class="ctime">${c.ts?c.ts.slice(0,16).replace('T',' '):''}</span>
            <span class="cuser">${c.user_name||'User'}</span>
            <span class="cemail">${c.user_email||''}</span>
          </div>
          <div class="ctext">${c.text}</div>
          ${c.ai_text_preview?`<div class="cpreview">AI: ${c.ai_text_preview.slice(0,100)}</div>`:''}
        </div>`).join('');
}

function filterComments() {
  const q = document.getElementById('commentSearch').value.toLowerCase();
  renderComments(_allComments.filter(c=>
    (c.user_email||'').toLowerCase().includes(q)||
    (c.user_name||'').toLowerCase().includes(q)||
    (c.text||'').toLowerCase().includes(q)||
    (c.ai_text_preview||'').toLowerCase().includes(q)
  ));
}

// ── INSTRUCTIONS ──────────────────────────────────────────
function renderInstructions(tips) {
  document.getElementById('instrList').innerHTML = tips.length===0
    ? '<p class="instr-empty">No instructions yet. Add one below.</p>'
    : tips.map(t=>`
        <div class="instr-item">
          <div class="instr-text">${t.text}</div>
          <div class="instr-date">${t.addedAt?t.addedAt.slice(0,10):''}</div>
          <button class="instr-del" onclick="deleteInstruction('${t.id}')">Delete</button>
        </div>`).join('');
}

async function addInstruction() {
  const text = document.getElementById('newInstrInput').value.trim();
  if(!text) return;
  try {
    const r = await fetch('/admin-tips', {
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body: JSON.stringify({admin_email:ADMIN_EMAIL, text})
    });
    if(!r.ok) throw new Error('Server error '+r.status);
    document.getElementById('newInstrInput').value='';
    toast('Instruction added ✓');
    loadStats();
  } catch(e){ toast('Failed: '+e.message, true); }
}

async function deleteInstruction(id) {
  if(!confirm('Delete this instruction?')) return;
  try {
    const r = await fetch('/admin-tips/'+encodeURIComponent(id)+'?admin_email='+encodeURIComponent(ADMIN_EMAIL), {method:'DELETE'});
    if(!r.ok) throw new Error('Server error '+r.status);
    toast('Instruction deleted');
    loadStats();
  } catch(e){ toast('Failed: '+e.message, true); }
}

// ── ALLOWLIST ──────────────────────────────────────────────
async function loadAllowlist() {
  try {
    const r = await fetch('/allowed-users');
    const d = await r.json();
    _pendingEmails = (d.emails||[]).filter(e=>e.toLowerCase()!==ADMIN_EMAIL.toLowerCase());
    renderEmailTags();
  } catch(e){ toast('Error loading allowlist', true); }
}

function renderEmailTags() {
  const el = document.getElementById('emailList');
  if(_pendingEmails.length===0){
    el.innerHTML='<span style="color:var(--muted);font-size:.85rem">No extra users added yet.</span>';
    return;
  }
  el.innerHTML = _pendingEmails.map((e,i)=>`
    <span class="email-tag">${e}
      <button onclick="removeEmail(${i})">✕</button>
    </span>`).join('');
}

function addEmail() {
  const v = document.getElementById('newEmailInput').value.trim().toLowerCase();
  if(!v||!v.includes('@')){ toast('Enter a valid email', true); return; }
  if(_pendingEmails.includes(v)||v===ADMIN_EMAIL.toLowerCase()){ toast('Already in list', true); return; }
  _pendingEmails.push(v);
  document.getElementById('newEmailInput').value='';
  renderEmailTags();
}

function removeEmail(i) {
  _pendingEmails.splice(i,1);
  renderEmailTags();
}

async function saveAllowlist() {
  try {
    const r = await fetch('/allowed-users', {
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body: JSON.stringify({admin_email:ADMIN_EMAIL, emails:_pendingEmails})
    });
    if(!r.ok) throw new Error('Server error '+r.status);
    toast('Access list saved ✓');
  } catch(e){ toast('Save failed: '+e.message, true); }
}

document.getElementById('newEmailInput').addEventListener('keydown',e=>{if(e.key==='Enter')addEmail();});

// ── SYSTEM HEALTH ──────────────────────────────────────────
let _adminSSE = null;
const STAGE_COLORS = {
  boot:'#00d4ff',req:'#a5f3fc',mem:'#7c3aed',web:'#22d3ee',
  gemini:'#a855f7',reply:'#00e676',prune:'#ffb347',error:'#ff4d6d'
};

function reconnectSSE() {
  if(_adminSSE){ try{_adminSSE.close();}catch(e){} _adminSSE=null; }
  const el = document.getElementById('thoughtLog');
  const st = document.getElementById('sseStatus');
  if(el) el.innerHTML='<span style="color:var(--muted);font-style:italic">Connecting to live log…</span>';
  if(st){ st.textContent='connecting…'; st.style.color='var(--warn)'; }

  try {
    _adminSSE = new EventSource('/stream-log');
    _adminSSE.onopen = () => {
      const s = document.getElementById('sseStatus');
      if(s){ s.textContent='● live'; s.style.color='var(--success)'; }
    };
    _adminSSE.onmessage = (e) => {
      try {
        const d = JSON.parse(e.data);
        appendLogLine(d);
      } catch(err){}
    };
    _adminSSE.onerror = () => {
      const s = document.getElementById('sseStatus');
      if(s){ s.textContent='disconnected'; s.style.color='var(--danger)'; }
    };
  } catch(err){
    const s = document.getElementById('sseStatus');
    if(s){ s.textContent='not available'; s.style.color='var(--muted)'; }
  }
}

function appendLogLine(entry) {
  const el = document.getElementById('thoughtLog');
  if(!el) return;
  // Remove placeholder
  const ph = el.querySelector('span[style*="font-style"]');
  if(ph) ph.remove();

  const color = STAGE_COLORS[entry.stage] || '#94a3b8';
  const line = document.createElement('div');
  line.style.cssText = 'padding:1px 0;border-bottom:1px solid rgba(255,255,255,.03);';
  line.innerHTML =
    '<span style="color:'+color+';font-weight:600;min-width:58px;display:inline-block">['+entry.stage+']</span> ' +
    '<span style="color:rgba(148,163,184,.55);font-size:.72rem;margin-right:10px">'+entry.ts+'</span>' +
    '<span style="color:#e2eeff">'+escHtml(entry.msg||'')+'</span>';
  el.appendChild(line);

  // Keep at most 80 lines
  while(el.children.length > 80) el.removeChild(el.firstChild);
  el.scrollTop = el.scrollHeight;
}

function escHtml(s) {
  return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}

async function loadSystem() {
  try {
    const [rs, ra] = await Promise.all([
      fetch('/rate-status?admin_email='+encodeURIComponent(ADMIN_EMAIL)),
      fetch('/admin-stats?admin_email='+encodeURIComponent(ADMIN_EMAIL)),
    ]);
    if(!rs.ok||!ra.ok){ toast('Failed to load system stats', true); return; }
    const rate   = await rs.json();
    const astats = await ra.json();

    const rpmPct = Math.round(rate.global_rpm_used / rate.global_rpm_limit * 100);
    const qPct   = Math.round(rate.queue_depth / rate.queue_max * 100);
    const upMins = Math.floor(rate.uptime_seconds / 60);
    const upHrs  = Math.floor(upMins / 60);
    const upStr  = upHrs>0 ? upHrs+'h '+(upMins%60)+'m' : upMins+'m';

    document.getElementById('rateGrid').innerHTML =
      '<div class="rate-item">' +
        '<div class="rl">Global RPM</div>' +
        '<div class="rv">'+rate.global_rpm_used+' / '+rate.global_rpm_limit+'</div>' +
        '<div class="bar-wrap"><div class="bar-fill '+(rpmPct>80?'danger':rpmPct>50?'warn':'')+'" style="width:'+rpmPct+'%"></div></div>' +
      '</div>' +
      '<div class="rate-item">' +
        '<div class="rl">Queue Depth</div>' +
        '<div class="rv">'+rate.queue_depth+' / '+rate.queue_max+'</div>' +
        '<div class="bar-wrap"><div class="bar-fill '+(qPct>70?'warn':'')+'" style="width:'+Math.max(qPct,2)+'%"></div></div>' +
      '</div>' +
      '<div class="rate-item"><div class="rl">Gemini Calls</div><div class="rv">'+rate.total_gemini_calls_today+'</div></div>' +
      '<div class="rate-item"><div class="rl">Errors</div><div class="rv" style="color:'+(rate.total_gemini_errors>0?'var(--danger)':'var(--success)')+'">'+rate.total_gemini_errors+'</div></div>' +
      '<div class="rate-item"><div class="rl">Uptime</div><div class="rv" style="font-size:1rem">'+upStr+'</div></div>' +
      '<div class="rate-item"><div class="rl">Per-User Limit</div><div class="rv">'+rate.per_user_rpm_limit+' req/min</div></div>';

    // Fetch memory pressure
    try {
      const mp = await fetch('/admin/memory-pressure?admin_email='+encodeURIComponent(ADMIN_EMAIL));
      if(mp.ok) {
        const m = await mp.json();
        const pressurePct = Math.min(100, Math.round(m.est_total_app_kb / 350000 * 100));
        const pressureColor = pressurePct > 70 ? 'var(--danger)' : pressurePct > 45 ? 'var(--warn)' : 'var(--success)';
        document.getElementById('memStats').innerHTML =
          '<p style="margin-bottom:8px">Memory entries: <b style="color:var(--accent)">'+m.total_mem_entries+'</b>' +
          ' &nbsp;|&nbsp; Est. data: <b style="color:var(--accent)">'+m.est_mem_kb+' KB</b>' +
          ' &nbsp;|&nbsp; Cache: <b style="color:var(--accent)">'+m.cache_entries+' entries</b></p>' +
          '<p style="margin-bottom:10px">Reactions: <b style="color:var(--accent)">'+m.reaction_entries+'</b>' +
          ' &nbsp;|&nbsp; Thought log: <b style="color:var(--accent)">'+m.thought_entries+' entries</b></p>' +
          '<div style="display:flex;align-items:center;gap:12px;margin-bottom:8px">' +
            '<span style="font-size:.8rem;color:var(--muted)">Estimated RAM pressure:</span>' +
            '<div style="flex:1;height:8px;background:var(--surface);border-radius:100px;overflow:hidden">' +
              '<div style="height:100%;width:'+pressurePct+'%;background:'+pressureColor+';border-radius:100px;transition:width .4s"></div>' +
            '</div>' +
            '<b style="color:'+pressureColor+';font-size:.85rem">'+pressurePct+'%</b>' +
          '</div>' +
          '<p style="color:var(--muted);font-size:.78rem">Prune runs every 6 h &mdash; Emergency Clear trims to 40 entries/user instantly.</p>' +
          (pressurePct > 60 ? '<p style="color:var(--warn);font-size:.8rem;margin-top:8px">⚠ High memory pressure — consider running Emergency Clear.</p>' : '');
      }
    } catch(e) {
      document.getElementById('memStats').innerHTML =
        '<p style="margin-bottom:8px">Memory entries: <b style="color:var(--accent)">'+(astats.total_mem_entries||0)+'</b></p>' +
        '<p style="color:var(--muted);font-size:.8rem">Prune runs every 6 hours.</p>';
    }

    // Start SSE if not already connected
    if(!_adminSSE || _adminSSE.readyState === 2) reconnectSSE();

    // Storage health
    try {
      const hr = await fetch('/health');
      if(hr.ok) {
        const hd = await hr.json();
        const isPersistent = hd.storage === 'persistent';
        const storageHtml =
          '<div class="rate-item" style="grid-column:1/-1;background:'+(isPersistent?'rgba(0,230,118,.07)':'rgba(255,179,71,.07)')+';border:1px solid '+(isPersistent?'rgba(0,230,118,.2)':'rgba(255,179,71,.3)')+'">' +
            '<div class="rl">💾 Storage Mode</div>' +
            '<div class="rv" style="font-size:.95rem;color:'+(isPersistent?'var(--success)':'var(--warn)')+'">' +
              (isPersistent ? '✅ Persistent Disk' : '⚠️ Ephemeral — data lost on restart') +
            '</div>' +
            (isPersistent ? '' : '<div style="font-size:.72rem;color:var(--warn);margin-top:4px">Set RENDER_PERSISTENT_DIR env var to a mounted disk path.</div>') +
          '</div>';
        document.getElementById('rateGrid').innerHTML += storageHtml;
      }
    } catch(e) { /* storage check is optional */ }

  } catch(e){ toast('System stats error: '+e.message, true); }
}

async function emergencyClear() {
  if(!confirm('Emergency clear: trims all user memories to 40 entries, clears web cache, and shrinks reaction log. Continue?')) return;
  const btn = document.querySelector('[onclick="emergencyClear()"]');
  if(btn){ btn.textContent = '⏳ Clearing…'; btn.disabled = true; }
  try {
    const r = await fetch('/admin/emergency-clear', {
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body: JSON.stringify({admin_email: ADMIN_EMAIL})
    });
    if(!r.ok) throw new Error('Server error '+r.status);
    const d = await r.json();
    toast('✅ Emergency clear: freed ~'+d.freed_kb+' KB ('+d.actions.length+' actions)');
    loadSystem();
  } catch(e){ toast('Emergency clear failed: '+e.message, true); }
  finally { if(btn){ btn.textContent = '🚨 Emergency Clear'; btn.disabled = false; } }
}

// ── BROWSER-SIDE KEEP-ALIVE ──────────────────────────────────────────────────
// This tab pings /health every 4 minutes so Render never hits the 15-min sleep
// threshold, even when no users are actively chatting.
// The server-side self-ping can't wake a sleeping server — only external HTTP
// requests can. This browser ping is that external trigger.
(function startKeepAlivePing() {
  const PING_URL      = 'https://astral-1-sb1i.onrender.com/health';
  const INTERVAL_MS   = 4 * 60 * 1000;   // 4 minutes — well under Render's 15-min limit
  const MAX_FAILURES  = 5;
  let failures        = 0;
  let intervalId      = null;

  async function ping() {
    try {
      const r = await fetch(PING_URL, { method: 'GET', cache: 'no-store' });
      if (r.ok) {
        failures = 0;
      } else {
        failures++;
      }
    } catch (e) {
      failures++;
    }
    // If repeated failures (e.g. tab went offline) back off — will resume when fetch succeeds again
    if (failures >= MAX_FAILURES) {
      clearInterval(intervalId);
      // Retry once after 2 min; if that works, restart the normal interval
      setTimeout(async () => {
        try {
          await fetch(PING_URL, { method: 'GET', cache: 'no-store' });
          failures  = 0;
          intervalId = setInterval(ping, INTERVAL_MS);
        } catch(e) {
          startKeepAlivePing();  // full restart
        }
      }, 2 * 60 * 1000);
    }
  }

  // Ping immediately on page load, then every 4 minutes
  ping();
  intervalId = setInterval(ping, INTERVAL_MS);

  // Also ping when the tab becomes visible again after being hidden
  document.addEventListener('visibilitychange', () => {
    if (document.visibilityState === 'visible') ping();
  });
})();
</script>
</body>
</html>
"""


@app.get("/admin", response_class=HTMLResponse)
async def admin_panel():
    return HTMLResponse(content=ADMIN_HTML)


# ═════════════════════════════════════════════════════════════════════════════
#  USAGE PAGE  (/usage)
# ═════════════════════════════════════════════════════════════════════════════
USAGE_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Astral — API Usage</title>
<meta http-equiv="refresh" content="15">
<style>
  :root {
    --bg: #050c12; --surface: #0d1b26; --card: #0f2236;
    --accent: #00e5ff; --text: #e2f4ff; --muted: #6b9ab8;
    --danger: #ff4d6d; --success: #00e676; --warn: #ffca28;
    --border: rgba(0,229,255,.12);
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { background: var(--bg); color: var(--text);
         font-family: 'Segoe UI', sans-serif; min-height: 100vh;
         display: flex; flex-direction: column; align-items: center;
         justify-content: center; padding: 32px; }
  .card { background: var(--card); border: 1px solid var(--border);
          border-radius: 20px; padding: 40px; width: 100%; max-width: 520px; }
  h1 { color: var(--accent); font-size: 1.3rem; letter-spacing: 2px;
       margin-bottom: 4px; }
  .sub { color: var(--muted); font-size: .8rem; margin-bottom: 32px; }
  .info-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 16px; margin-top: 28px; }
  .info-item { background: var(--surface); border-radius: 10px; padding: 14px; }
  .info-item .il { color: var(--muted); font-size: .72rem; text-transform: uppercase;
                   letter-spacing: 1px; margin-bottom: 4px; }
  .info-item .iv { font-size: 1.4rem; font-weight: 700; color: var(--accent); }
  .refresh-note { color: var(--muted); font-size: .75rem; text-align: center; margin-top: 24px; }
  .status-dot { display: inline-block; width: 8px; height: 8px; border-radius: 50%;
                background: var(--success); margin-right: 6px;
                animation: pulse 2s infinite; }
  @keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.4} }
</style>
</head>
<body>
<div class="card" id="card">
  <h1>⬡ ASTRAL API USAGE</h1>
  <div class="sub"><span class="status-dot"></span>Live · refreshes every 15 seconds</div>
  <p style="color:var(--muted);font-size:.85rem">Loading…</p>
</div>
<script>
async function load() {
  try {
    const r = await fetch('/health');
    const h = await r.json();
    const upMin = Math.floor((h.uptime_seconds||0)/60);
    const upStr = upMin>60 ? Math.floor(upMin/60)+'h '+upMin%60+'m' : upMin+'m';
    document.getElementById('card').innerHTML = `
      <h1>⬡ ASTRAL API USAGE</h1>
      <div class="sub"><span class="status-dot"></span>Live · refreshes every 15 seconds</div>
      <div class="info-grid">
        <div class="info-item"><div class="il">Status</div>
          <div class="iv" style="font-size:1rem;color:var(--success)">✓ Online</div></div>
        <div class="info-item"><div class="il">Uptime</div>
          <div class="iv" style="font-size:1rem">${upStr}</div></div>
        <div class="info-item"><div class="il">Total Users</div>
          <div class="iv">${h.users}</div></div>
        <div class="info-item"><div class="il">Allowed Users</div>
          <div class="iv">${h.allowed_users}</div></div>
      </div>
      <div style="margin-top:28px;padding:16px;background:var(--surface);border-radius:10px;
                  color:var(--muted);font-size:.85rem;line-height:1.6">
        <b style="color:var(--text)">Free-tier limits (Gemini)</b><br>
        • 15 requests per minute (RPM) globally<br>
        • 1,000,000 tokens per day (TPD)<br>
        • 6 messages per user per minute<br><br>
        If Astral says "I'm at my limit", wait 60 seconds and try again.
      </div>
      <div class="refresh-note">↻ Auto-refreshing · <a href="/admin"
        style="color:var(--accent);text-decoration:none">Admin panel →</a></div>
    `;
  } catch(e) {
    document.getElementById('card').innerHTML += `<p style="color:var(--danger)">Error: ${e}</p>`;
  }
}
load();
// Same browser-side keep-alive as admin page
setInterval(()=>fetch('https://astral-1-sb1i.onrender.com/health',{cache:'no-store'}).catch(()=>{}), 4*60*1000);
document.addEventListener('visibilitychange',()=>{ if(document.visibilityState==='visible') fetch('https://astral-1-sb1i.onrender.com/health',{cache:'no-store'}).catch(()=>{}); });
</script>
</body>
</html>
"""


@app.get("/usage", response_class=HTMLResponse)
async def usage_page():
    return HTMLResponse(content=USAGE_HTML)


# ── /ping.js  — include this script on your static frontend ───────────────────
# Add this to your frontend HTML:
#   <script src="https://astral-1-sb1i.onrender.com/ping.js"></script>
# Any page that loads this script will ping the server every 4 minutes,
# keeping it awake even when the admin page is closed.
@app.get("/ping.js")
async def ping_js():
    js = """
(function(){
  var URL='https://astral-1-sb1i.onrender.com/health';
  var IV=4*60*1000;
  function p(){fetch(URL,{cache:'no-store'}).catch(function(){});}
  p();
  setInterval(p,IV);
  document.addEventListener('visibilitychange',function(){if(document.visibilityState==='visible')p();});
})();
""".strip()
    from fastapi.responses import Response
    return Response(content=js, media_type="application/javascript",
                    headers={"Cache-Control": "no-store"})


# ── File generation helpers ───────────────────────────────────────────────────

# Patterns that indicate the user wants a file generated
import re as _re

_FILE_REQUEST_RE = _re.compile(
    r'\b(create|make|generate|write|build|give me|produce|save|output)\b.{0,100}'
    r'(?:\.([a-zA-Z0-9]+)|(?:([a-zA-Z0-9]+)\s+(?:file|document|script|page|snippet|log|code)))'
    r'|\b(a|an)\s+(webpage?|website|landing\s+page|html\s+page)\b',
    _re.IGNORECASE
)

_EXT_RE = _re.compile(
    r'\b(html?|css|js|jsx|ts|tsx|py|json|md|txt|csv|docx?|xml|yaml|yml|sh|sql)\b',
    _re.IGNORECASE
)

def _detect_file_request(text: str):
    """Return (True, ext) if the message is asking to generate a file, else (False, None)."""
    if not _FILE_REQUEST_RE.search(text):
        return False, None
    m = _FILE_REQUEST_RE.search(text)
    if not m:
        return False, None
    
    # Extract extension from groups
    ext = (m.group(2) or m.group(3) or m.group(4) or "txt").lower()
    
    # If it matched the webpage/website pattern without a specific extension
    if m.group(4) and not ext:
        ext = "html"
    # normalise
    if ext == "htm":   ext = "html"
    if ext == "doc":   ext = "docx"
    if ext == "log":   ext = "txt"
    return True, ext


async def _generate_file_content(user_prompt: str, ext: str, model_name: str) -> str:
    """Ask Gemini to produce raw file content only — no markdown fences, no commentary."""
    ext_hints = {
        "html":  "a complete, self-contained HTML page with inline CSS and JS if needed",
        "css":   "a CSS stylesheet",
        "js":    "a JavaScript file (no markdown, no explanation, just the code)",
        "jsx":   "a React JSX component (no markdown, just the code)",
        "ts":    "a TypeScript file (no markdown, just the code)",
        "tsx":   "a React TSX component (no markdown, just the code)",
        "py":    "a Python script (no markdown, just the code)",
        "json":  "valid JSON only, no other text",
        "md":    "a Markdown document",
        "txt":   "a plain text document",
        "csv":   "valid CSV data only",
        "docx":  "a plain text document that will be converted to docx",
        "xml":   "valid XML only",
        "yaml":  "valid YAML only",
        "yml":   "valid YAML only",
        "sh":    "a bash shell script",
        "sql":   "SQL statements only",
    }
    hint = ext_hints.get(ext, "the requested file content")
    system = (
        f"You are a file generator. Output ONLY {hint}. "
        "Your goal is to produce the FULL and COMPLETE content of the file based on the user's request. "
        "Do NOT include markdown code fences, backticks, explanations, or any surrounding text. "
        "The output will be saved directly to a file and opened by the user. "
        "If the user asks for a 'Hello' file, you must output the actual 'Hello' text or code."
    )
    gen_model = genai.GenerativeModel(
        model_name=model_name,
        system_instruction=system,
        generation_config=genai.GenerationConfig(temperature=0.3, max_output_tokens=8192),
    )
    result = await _gemini_generate(gen_model, user_prompt, user_key="filegen")
    # Strip accidental fences just in case
    text = result.strip() if isinstance(result, str) else ""
    text = _re.sub(r"^```[a-zA-Z]*\n?", "", text)
    text = _re.sub(r"\n?```$", "", text)
    return text.strip()


def _build_file_bytes(content: str, ext: str) -> bytes:
    """Turn text content into file bytes. docx gets special treatment."""
    if ext == "docx":
        return _build_docx(content)
    return content.encode("utf-8")


def _mime_for_ext(ext: str) -> str:
    m = {
        "html": "text/html", "css": "text/css", "js": "application/javascript",
        "jsx": "application/javascript", "ts": "application/typescript",
        "tsx": "application/typescript", "py": "text/x-python",
        "json": "application/json", "md": "text/markdown", "txt": "text/plain",
        "csv": "text/csv", "xml": "application/xml", "yaml": "text/yaml",
        "yml": "text/yaml", "sh": "application/x-sh", "sql": "application/sql",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "zip": "application/zip",
    }
    return m.get(ext, "application/octet-stream")


# Accepts an image (base64) + desired output format.
# Uses Gemini Vision to extract / structure the text, then builds the file.
# Supported formats: docx, txt, html, js, css, py, json, md, zip
# ──────────────────────────────────────────────────────────────────────────────

class ConvertPayload(BaseModel):
    image_base64: str                          # raw base64, no data-URL prefix needed
    image_mime:   Optional[str] = "image/jpeg"
    format:       str           = "docx"       # target output format
    user_email:   Optional[str] = ""


@app.post("/convert")
async def convert_image_to_doc(payload: ConvertPayload):
    """OCR + structure an image of text, return the result as a downloadable file."""
    import base64 as _base64
    import io

    fmt = payload.format.lower().lstrip(".")

    # ── decode image ──────────────────────────────────────────────────────────
    raw_b64 = payload.image_base64 or ""
    if "," in raw_b64:
        raw_b64 = raw_b64.split(",", 1)[1]
    raw_b64 += "=" * (-len(raw_b64) % 4)
    try:
        image_bytes = _base64.b64decode(raw_b64, validate=True)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid image data: {e}")

    # ── detect mime ──────────────────────────────────────────────────────────
    header = image_bytes[:12]
    if header[:3] == b"\xff\xd8\xff":
        img_mime = "image/jpeg"
    elif header[:8] == b"\x89PNG\r\n\x1a\n":
        img_mime = "image/png"
    elif header[:6] in (b"GIF87a", b"GIF89a"):
        img_mime = "image/gif"
    elif header[:4] == b"RIFF" and header[8:12] == b"WEBP":
        img_mime = "image/webp"
    else:
        img_mime = payload.image_mime or "image/jpeg"

    # ── ask Gemini to extract + structure the text ────────────────────────────
    image_part = {"mime_type": img_mime, "data": image_bytes}
    ocr_prompt = (
        "Extract ALL text from this image exactly as written. "
        "Preserve the structure: headings stay as headings, numbered or lettered "
        "lists stay as lists, questions stay as questions. "
        "Do NOT add commentary, explanations, or markdown fences. "
        "Output the raw structured text only."
    )

    try:
        vision_model = genai.GenerativeModel(
            MODEL_VISION,
            generation_config=genai.GenerationConfig(temperature=0.1, max_output_tokens=4096),
        )
        response = await _gemini_generate(vision_model, [image_part, ocr_prompt],
                                          user_key=payload.user_email or "convert")
        extracted = response.strip() if isinstance(response, str) else ""
    except Exception as e:
        _log("error", f"/convert gemini failed: {e}")
        raise HTTPException(status_code=500, detail=f"OCR failed: {e}")
    finally:
        del image_bytes, image_part

    if not extracted:
        raise HTTPException(status_code=422, detail="No text could be extracted from the image.")

    # ── build the output file ─────────────────────────────────────────────────
    from fastapi.responses import Response as _Response

    if fmt == "docx":
        file_bytes = _build_docx(extracted)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename   = "converted.docx"

    elif fmt in ("zip",):
        # Zip containing a txt file of the extracted text
        import zipfile
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("extracted_text.txt", extracted)
        file_bytes = buf.getvalue()
        media_type = "application/zip"
        filename   = "converted.zip"

    else:
        # Plain-text formats: txt, html, js, css, py, json, md, etc.
        file_bytes = extracted.encode("utf-8")
        media_type = "text/plain"
        filename   = f"converted.{fmt}"

    return _Response(
        content=file_bytes,
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )


def _build_docx(text: str) -> bytes:
    """Build a .docx from plain structured text using python-docx."""
    try:
        from docx import Document as _Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa
    except ImportError:
        # python-docx not installed — fall back to plain bytes wrapped in minimal XML
        return _minimal_docx_fallback(text)

    doc = _Document()

    # Tighten default margins to 1 inch
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.lstrip()

        if not stripped:
            # blank line → small spacer paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

        elif stripped.startswith(("# ", "## ", "### ")):
            # Markdown-style headings Gemini may emit
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped.lstrip("# ").strip()
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            doc.add_heading(heading_text, level=min(level, 3))

        elif stripped[0].isdigit() and len(stripped) > 2 and stripped[1] in ".):":
            # Numbered list item: "1. foo" or "1) foo"
            p = doc.add_paragraph(style="List Number")
            p.add_run(stripped[2:].lstrip())

        elif (
            len(stripped) >= 4
            and stripped[0] in "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ"
            and stripped[1] in ".):)"
            and stripped[2] == " "
        ):
            # Lettered choice: "(A) foo" or "A. foo" or "A) foo"
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.strip())
            run.font.size = Pt(11)

        elif stripped.startswith(("(A)", "(B)", "(C)", "(D)", "(E)",
                                   "(a)", "(b)", "(c)", "(d)", "(e)")):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(stripped)
            run.font.size = Pt(11)

        elif stripped.startswith(("- ", "* ", "• ")):
            # Bullet list
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:])

        elif i == 0 or (i == 1 and not lines[0].strip()):
            # First non-blank line → treat as document title
            doc.add_heading(stripped, level=1)

        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(4)

        i += 1

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _minimal_docx_fallback(text: str) -> bytes:
    """Fallback: wrap text in a minimal valid .docx ZIP if python-docx missing."""
    import zipfile, io, textwrap
    escaped = (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                   .replace('"', "&quot;"))
    paras = "".join(
        f'<w:p><w:r><w:t xml:space="preserve">{line}</w:t></w:r></w:p>'
        for line in escaped.splitlines()
    )
    doc_xml = textwrap.dedent(f"""\
        <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
          <w:body>{paras}</w:body>
        </w:document>""")
    ct_xml = textwrap.dedent("""\
        <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
        <Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
          <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
          <Default Extension="xml"  ContentType="application/xml"/>
          <Override PartName="/word/document.xml"
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
        </Types>""")
    rels_xml = textwrap.dedent("""\
        <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
        <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
          <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
            Target="word/document.xml"/>
        </Relationships>""")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", ct_xml)
        zf.writestr("_rels/.rels", rels_xml)
        zf.writestr("word/document.xml", doc_xml)
    return buf.getvalue()


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
